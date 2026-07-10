from __future__ import annotations

import json
import re
import urllib.error
import urllib.parse
import urllib.request
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from .errors import ValidationError

MAX_DOWNLOAD_BYTES = 25 * 1024 * 1024
CHUNK_TOKENS = 800
CHUNK_OVERLAP_TOKENS = 125
CHUNK_TOKEN_SAFETY_MARGIN = 75
MIN_CHUNK_TOKENS = 120

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".csv", ".tsv", ".json"}
WORD_EXTENSIONS = {".docx"}
TEXT_CONTENT_TYPES = {
    "text/plain",
    "text/markdown",
    "text/csv",
    "text/tab-separated-values",
    "application/json",
}
WORD_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
IMAGE_CONTENT_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/webp", "image/gif"}


def prepare_attachments(
    *,
    research_id: str,
    job_dir: Path,
    attachments: list[dict[str, Any]],
) -> dict[str, Any]:
    if not isinstance(attachments, list):
        raise ValidationError("attachments must be an array")
    files_dir = job_dir / "attachment-files"
    files_dir.mkdir(parents=True, exist_ok=True)

    context_files: list[dict[str, Any]] = []
    chunks: list[dict[str, Any]] = []
    for index, attachment in enumerate(attachments, start=1):
        if not isinstance(attachment, dict):
            context_files.append(_failed_file(index, {}, "attachment must be an object"))
            continue
        file_id = f"file-{index}"
        name = str(attachment.get("name") or Path(str(attachment.get("path") or f"attachment-{index}")).name).strip()
        path = str(attachment.get("path") or "").strip()
        content_type = str(attachment.get("content_type") or "").strip()
        size_bytes = _optional_int(attachment.get("size_bytes"))
        file_view = {
            "id": file_id,
            "name": name or f"attachment-{index}",
            "path": path,
            "content_type": content_type,
            "size_bytes": size_bytes,
        }
        try:
            download_url = str(attachment.get("download_url") or "").strip()
            if not download_url:
                raise ValidationError("download_url is required")
            local_path = files_dir / f"{file_id}-{_safe_filename(name or 'attachment')}"
            download_attachment(download_url, local_path)
            local_path_view = _relative_path(local_path, job_dir)
            if _is_image_attachment(name=name, path=path, content_type=content_type):
                if attachment.get("image_analysis_error"):
                    raise ValidationError(f"image analysis failed: {attachment.get('image_analysis_error')}")
                image_analysis = _validated_image_analysis(attachment.get("image_analysis"))
                image_summary = str(image_analysis.get("summary") or "").strip()
                image_relevance = _image_relevance_text(image_analysis)
                image_relevance_score = _image_relevance_score(image_analysis)
                context_files.append(
                    {
                        **file_view,
                        "text_chars": len(image_summary),
                        "chunk_count": 0,
                        "status": "ready",
                        "error": None,
                        "local_path": local_path_view,
                        "analysis": {
                            "type": "image",
                            "source": "analyze_image",
                            "summary": image_summary,
                            "key_points": [],
                            "relevance": image_relevance,
                            "relevance_score": image_relevance_score,
                            "selected_chunk_ids": [],
                            "payload": image_analysis,
                        },
                    }
                )
                continue
            text = extract_text(local_path, name=name, content_type=content_type)
            text = _clean_text(text)
            file_chunks = chunk_text(text)
            for chunk_index, chunk in enumerate(file_chunks, start=1):
                chunks.append(
                    {
                        "chunk_id": f"{file_id}:{chunk_index:04d}",
                        "file_id": file_id,
                        "file_name": file_view["name"],
                        "path": path,
                        "content_type": content_type,
                        "index": chunk_index,
                        "text": chunk,
                    }
                )
            context_files.append(
                {
                    **file_view,
                    "text_chars": len(text),
                    "chunk_count": len(file_chunks),
                    "status": "ready",
                    "error": None,
                    "local_path": local_path_view,
                }
            )
        except Exception as exc:  # noqa: BLE001
            context_files.append({**file_view, "text_chars": 0, "chunk_count": 0, "status": "failed", "error": str(exc)})

    return {
        "version": 1,
        "research_id": research_id,
        "prepared_at": datetime.now(timezone.utc).isoformat(timespec="milliseconds").replace("+00:00", "Z"),
        "files": context_files,
        "chunks": chunks,
        "summary": build_attachment_summary(context_files, chunks),
    }


def download_attachment(url: str, destination: Path) -> None:
    parsed = urllib.parse.urlparse(url)
    if parsed.scheme not in {"http", "https", "data"}:
        raise ValidationError("unsupported attachment download URL scheme")
    destination.parent.mkdir(parents=True, exist_ok=True)
    total = 0
    try:
        with urllib.request.urlopen(url, timeout=30) as response, destination.open("wb") as output:
            while True:
                chunk = response.read(64 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > MAX_DOWNLOAD_BYTES:
                    raise ValidationError("attachment exceeds the 25 MB download limit")
                output.write(chunk)
    except urllib.error.URLError as exc:
        raise ValidationError(f"attachment download failed: {exc}") from exc


def extract_text(path: Path, *, name: str, content_type: str = "") -> str:
    suffix = Path(name or path.name).suffix.lower() or path.suffix.lower()
    normalized_type = content_type.split(";")[0].strip().lower()
    if suffix in TEXT_EXTENSIONS or normalized_type in TEXT_CONTENT_TYPES or normalized_type.startswith("text/"):
        return _read_text(path)
    if suffix == ".pdf" or normalized_type == "application/pdf":
        return _extract_pdf_text(path)
    if suffix in WORD_EXTENSIONS or normalized_type in WORD_CONTENT_TYPES:
        return _extract_docx_text(path)
    raise ValidationError(f"unsupported attachment type: {content_type or suffix or 'unknown'}")


def _is_image_attachment(*, name: str, path: str, content_type: str) -> bool:
    suffix = Path(name or path).suffix.lower()
    normalized_type = content_type.split(";")[0].strip().lower()
    return suffix in IMAGE_EXTENSIONS or normalized_type in IMAGE_CONTENT_TYPES or normalized_type.startswith("image/")


def _validated_image_analysis(value: Any) -> dict[str, Any]:
    if not isinstance(value, dict):
        raise ValidationError("image_analysis is required for image attachments")
    summary = str(value.get("summary") or "").strip()
    if not summary:
        raise ValidationError("image_analysis.summary is required")
    return value


def _format_structured_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        parts = []
        for key, item in value.items():
            text = _format_structured_value(item)
            if text:
                parts.append(f"{key}: {text}")
        return "; ".join(parts)
    if isinstance(value, list):
        return "; ".join(text for text in (_format_structured_value(item) for item in value) if text)
    return str(value).strip()


def _image_relevance_text(value: Any) -> str:
    if not isinstance(value, dict):
        return ""
    return _format_structured_value(value.get("research_relevance"))


def _image_relevance_score(value: Any) -> float:
    if not isinstance(value, dict):
        return 0.0
    relevance = value.get("research_relevance")
    if isinstance(relevance, dict):
        explicit = _optional_float(relevance.get("relevance_score"))
        if explicit is not None:
            return explicit
        text = str(relevance.get("relevance") or "").strip().lower()
        if any(word in text for word in ("useful", "relevant", "supports", "related", "directly relevant")):
            return 0.8
        return 0.0
    text = str(relevance or "").strip().lower()
    if any(word in text for word in ("useful", "relevant", "supports", "related")):
        return 0.8
    return 0.0


def chunk_text(text: str) -> list[str]:
    clean = text.strip()
    if not clean:
        raise ValidationError("attachment did not contain extractable text")
    blocks = _split_structural_blocks(clean)
    chunks: list[str] = []
    current: list[str] = []
    current_tokens = 0
    previous_tail = ""

    for block in blocks:
        block_tokens = estimate_tokens(block)
        if block_tokens > CHUNK_TOKENS:
            if current:
                chunks.append("\n\n".join(current).strip())
                previous_tail = _chunk_tail(chunks[-1])
                current = []
                current_tokens = 0
            for piece in _split_long_block(block):
                if previous_tail:
                    piece = f"{previous_tail}\n\n{piece}"
                chunks.append(piece.strip())
                previous_tail = _chunk_tail(chunks[-1])
            continue

        if current and current_tokens + block_tokens > CHUNK_TOKENS:
            chunks.append("\n\n".join(current).strip())
            previous_tail = _chunk_tail(chunks[-1])
            if previous_tail and estimate_tokens(previous_tail) + block_tokens <= CHUNK_TOKENS:
                current = [previous_tail, block]
            else:
                current = [block]
            current_tokens = estimate_tokens("\n\n".join(current))
            continue

        current.append(block)
        current_tokens += block_tokens

    if current:
        chunks.append("\n\n".join(current).strip())

    return _merge_short_chunks([chunk for chunk in chunks if chunk.strip()])


def build_attachment_summary(files: list[dict[str, Any]], chunks: list[dict[str, Any]]) -> str:
    ready = [file for file in files if file.get("status") == "ready"]
    failed = [file for file in files if file.get("status") == "failed"]
    parts = [
        f"{len(ready)} uploaded file(s) prepared into {len(chunks)} text chunk(s).",
    ]
    if ready:
        parts.append(
            "Ready files: "
            + "; ".join(
                f"{file.get('name')} ({int(file.get('text_chars') or 0)} chars, {int(file.get('chunk_count') or 0)} chunks)"
                for file in ready[:8]
            )
        )
    if failed:
        parts.append("Failed files: " + "; ".join(f"{file.get('name')}: {file.get('error')}" for file in failed[:5]))
    preview = "\n\n".join(str(chunk.get("text") or "")[:500] for chunk in chunks[:3]).strip()
    if preview:
        parts.append("Preview:\n" + preview)
    return "\n".join(parts)


def _read_text(path: Path) -> str:
    data = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = data.decode("utf-8", errors="replace")
    if path.suffix.lower() == ".json":
        try:
            parsed = json.loads(text)
            return json.dumps(parsed, ensure_ascii=False, indent=2)
        except json.JSONDecodeError:
            return text
    return text


def _extract_pdf_text(path: Path) -> str:
    try:
        import fitz  # type: ignore[import-not-found]
    except Exception as exc:  # noqa: BLE001
        raise ValidationError("PDF attachment parsing requires PyMuPDF/fitz") from exc
    pages: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            pages.append(page.get_text("text"))
    return "\n\n".join(pages)


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore[import-not-found]
    except Exception as exc:  # noqa: BLE001
        raise ValidationError("DOCX attachment parsing requires python-docx") from exc
    try:
        document = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        raise ValidationError(f"failed to open DOCX attachment: {exc}") from exc

    blocks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)
    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [_docx_cell_text(cell) for cell in row.cells]
            clean_cells = [cell for cell in cells if cell]
            if clean_cells:
                rows.append(" | ".join(clean_cells))
        if rows:
            blocks.append("\n".join(rows))
    text = "\n\n".join(blocks).strip()
    if not text:
        raise ValidationError("DOCX attachment did not contain extractable text")
    return text


def _docx_cell_text(cell: Any) -> str:
    lines = [paragraph.text.strip() for paragraph in cell.paragraphs if paragraph.text.strip()]
    return " ".join(lines).strip()


def _clean_text(text: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", text.replace("\r\n", "\n").replace("\r", "\n")).strip()


def _split_structural_blocks(text: str) -> list[str]:
    blocks: list[str] = []
    current_heading = ""
    for raw in re.split(r"\n\s*\n", text):
        block = raw.strip()
        if not block:
            continue
        lines = [line.strip() for line in block.splitlines() if line.strip()]
        if len(lines) == 1 and _looks_like_heading(lines[0]):
            if current_heading:
                blocks.append(current_heading)
            current_heading = lines[0]
            continue
        if current_heading:
            block = f"{current_heading}\n{block}"
            current_heading = ""
        blocks.append(block)
    if current_heading:
        blocks.append(current_heading)
    return blocks or [text]


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if stripped.startswith(("#", "##", "###", "####")):
        return True
    if re.match(r"^(\d+(\.\d+)*[.)]?|[一二三四五六七八九十]+[、.．])\s*\S+", stripped):
        return True
    if len(stripped) <= 80 and not re.search(r"[。.!?！？；;]$", stripped):
        return True
    return False


def estimate_tokens(text: str) -> int:
    """Approximate mixed Chinese/English token count without adding tokenizer deps."""
    cjk = re.findall(r"[\u3400-\u9fff]", text)
    words = re.findall(r"[A-Za-z0-9]+(?:[-_./][A-Za-z0-9]+)*", text)
    punctuation = re.findall(r"[^\s\w\u3400-\u9fff]", text)
    return max(1, len(cjk) + len(words) + max(0, len(punctuation) // 4))


def _split_long_block(block: str) -> list[str]:
    units = _sentence_units(block)
    pieces: list[str] = []
    current: list[str] = []
    current_tokens = 0
    target_tokens = max(1, CHUNK_TOKENS - CHUNK_OVERLAP_TOKENS - CHUNK_TOKEN_SAFETY_MARGIN)
    for unit in units:
        unit_tokens = estimate_tokens(unit)
        if unit_tokens > target_tokens:
            if current:
                pieces.append("".join(current).strip())
                current = []
                current_tokens = 0
            pieces.extend(_split_oversized_unit(unit))
            continue
        if current and current_tokens + unit_tokens > target_tokens:
            pieces.append("".join(current).strip())
            current = []
            current_tokens = 0
        current.append(unit)
        current_tokens += unit_tokens
    if current:
        pieces.append("".join(current).strip())
    return [piece for piece in pieces if piece]


def _sentence_units(text: str) -> list[str]:
    units = re.findall(r".+?(?:[。！？.!?](?:\s+|$)|\n+|$)", text, flags=re.S)
    return [unit for unit in units if unit.strip()]


def _split_oversized_unit(text: str) -> list[str]:
    tokens = re.findall(r"[\u3400-\u9fff]|[A-Za-z0-9]+(?:[-_./][A-Za-z0-9]+)*|\S", text)
    pieces: list[str] = []
    start = 0
    target_tokens = max(1, CHUNK_TOKENS - CHUNK_OVERLAP_TOKENS - CHUNK_TOKEN_SAFETY_MARGIN)
    while start < len(tokens):
        end = min(len(tokens), start + target_tokens)
        piece = _join_tokenish(tokens[start:end]).strip()
        if piece:
            pieces.append(piece)
        if end >= len(tokens):
            break
        start = max(end - CHUNK_OVERLAP_TOKENS, start + 1)
    return pieces


def _join_tokenish(tokens: list[str]) -> str:
    text = ""
    for token in tokens:
        if not text:
            text = token
        elif re.match(r"^[A-Za-z0-9]", token) and re.search(r"[A-Za-z0-9]$", text):
            text += " " + token
        else:
            text += token
    return text


def _chunk_tail(text: str) -> str:
    units = _sentence_units(text)
    tail: list[str] = []
    total = 0
    for unit in reversed(units):
        unit_tokens = estimate_tokens(unit)
        if tail and total + unit_tokens > CHUNK_OVERLAP_TOKENS:
            break
        tail.insert(0, unit.strip())
        total += unit_tokens
        if total >= CHUNK_OVERLAP_TOKENS:
            break
    return "\n".join(tail).strip()


def _merge_short_chunks(chunks: list[str]) -> list[str]:
    if len(chunks) <= 1:
        return chunks
    merged: list[str] = []
    for chunk in chunks:
        if merged and estimate_tokens(chunk) < MIN_CHUNK_TOKENS and estimate_tokens(merged[-1]) + estimate_tokens(chunk) <= CHUNK_TOKENS:
            merged[-1] = f"{merged[-1]}\n\n{chunk}".strip()
        else:
            merged.append(chunk)
    return merged


def _safe_filename(name: str) -> str:
    clean = re.sub(r"[^\w.\-]+", "-", name.strip(), flags=re.UNICODE).strip("-._")
    return clean[:120] or "attachment"


def _optional_int(value: Any) -> int | None:
    try:
        if value is None or value == "":
            return None
        return int(value)
    except (TypeError, ValueError):
        return None


def _optional_float(value: Any) -> float | None:
    try:
        if value is None or value == "":
            return None
        number = float(value)
    except (TypeError, ValueError):
        return None
    return max(0.0, min(1.0, number))


def _relative_path(path: Path, root: Path) -> str:
    try:
        return path.relative_to(root).as_posix()
    except ValueError:
        return path.as_posix()


def _failed_file(index: int, attachment: dict[str, Any], error: str) -> dict[str, Any]:
    return {
        "id": f"file-{index}",
        "name": str(attachment.get("name") or f"attachment-{index}"),
        "path": str(attachment.get("path") or ""),
        "content_type": str(attachment.get("content_type") or ""),
        "size_bytes": _optional_int(attachment.get("size_bytes")),
        "text_chars": 0,
        "chunk_count": 0,
        "status": "failed",
        "error": error,
    }
