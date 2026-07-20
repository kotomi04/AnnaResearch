from __future__ import annotations

import io
import hashlib
import json
import math
import os
import asyncio
import threading
import time
import urllib.error
import urllib.request
import uuid

import pytest

from researcher_tool.attachments import chunk_text, extract_text
from researcher_tool import attachment_embeddings as attachment_embeddings_module
from researcher_tool.context_selector import LexicalContextSelector
from researcher_tool.hybrid_context_selector import HybridContextSelector, split_text_ranges
from researcher_tool.attachment_embeddings import embed_attachment_chunks
from researcher_tool.attachment_summary import select_attachment_context, summarize_attachment_context
from researcher_tool.dispatcher import AppDispatcher
from researcher_tool.embedding import MAX_PARALLEL_EMBEDDING_BATCHES, AnnaEmbeddingsClient, EmbeddingBatchOutcome, EmbeddingsError
from researcher_tool.errors import NotFoundError, ValidationError
from researcher_tool.job_store import JobStore, normalize_query_for_dedup
from researcher_tool.outline_discovery import generate_outline_draft
from researcher_tool.settings import SettingsStore, mask_secret
from researcher_tool.sources import (
    BUILTIN_SOURCE_IDS,
    CredentialStore,
    ResearchSourceExecutor,
    ResearchSourceRegistry,
    SourceCallError,
    builtin_duckduckgo_definition,
    builtin_tavily_definition,
    migrate_legacy_tavily_key,
)
from researcher_tool.sources.envelope import EnvelopeError, validate_envelope
from researcher_tool.sources.executor import resolve_path
from researcher_tool.sources.extraction import arxiv as arxiv_extraction
from researcher_tool.sources.extraction import browser_fallback
from researcher_tool.sources.extraction import fetcher as extraction_fetcher
from researcher_tool.sources.extraction import pdf as pdf_extraction
from researcher_tool.sources.extraction.html import extract_html
from researcher_tool.sources.extraction.models import ExtractedPage
from researcher_tool.sources.extraction.pdf import extract_pdf
from researcher_tool.sources.extraction.tavily import TAVILY_PREFETCH_MIN_CHARS, enrich_tavily_items
from researcher_tool.sources.extraction.utils import same_url_without_fragment
from researcher_tool.sources.native import duckduckgo as duckduckgo_native
from researcher_tool.sources.native.executor import NativeResearchSourceExecutor
from researcher_tool.views import compact_job_view
from researcher_tool.web_documents import WebDocumentStore


def make_dispatcher(tmp_path):
    root = tmp_path / ".research"
    return AppDispatcher(
        settings=SettingsStore(root=root),
        jobs=JobStore(root=root),
        selector=LexicalContextSelector(max_sources=4, context_budget=4000),
        transfers=MemoryTransfers(),
    )


def post_json(url: str, payload: dict):
    data = json.dumps(payload).encode("utf-8")
    request = urllib.request.Request(url, data=data, method="POST", headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(request, timeout=5) as response:
        return json.loads(response.read().decode("utf-8"))


def get_json(url: str):
    request = urllib.request.Request(url, method="GET")
    with urllib.request.urlopen(request, timeout=5) as response:
        return json.loads(response.read().decode("utf-8"))


class MemoryTransfers:
    def __init__(self):
        self.payloads = {}
        self.deleted = []

    def upload(self, *, prefix, kind, payload):
        data = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
        path = f"{prefix}/transfers/{kind}-{uuid.uuid4().hex}.json"
        self.payloads[path] = payload
        return {"path": path, "content_type": "application/json", "size_bytes": len(data), "sha256": hashlib.sha256(data).hexdigest(), "delete_after_read": True}

    def download_json(self, descriptor, *, expected_prefix):
        assert descriptor["path"].startswith(expected_prefix + "/transfers/")
        return self.payloads[descriptor["path"]]

    def delete_best_effort(self, descriptor):
        path = descriptor.get("path")
        self.deleted.append(path)
        self.payloads.pop(path, None)


class FakeEmbeddings:
    def __init__(self):
        self.calls = []
        self.batch_calls = []

    def create(self, *, texts, model="anna-managed-v1", timeout=30.0):
        self.calls.append(list(texts))
        return {
            "data": [{"embedding": [1.0, float(index + 1)]} for index, _ in enumerate(texts)],
            "_meta": {"dimensions": 2},
        }

    def create_batches(self, *, batches, model="anna-managed-v1", timeout=30.0):
        self.batch_calls.append([list(batch) for batch in batches])
        return [self.create(texts=batch, model=model, timeout=timeout) for batch in batches]

    def create_batches_settled(self, *, batches, model="anna-managed-v1", timeout=30.0):
        self.batch_calls.append([list(batch) for batch in batches])
        return [EmbeddingBatchOutcome(result=self.create(texts=batch, model=model, timeout=timeout)) for batch in batches]


class FakeSampling:
    def __init__(self):
        self.calls = []

    def create_message(self, **kwargs):
        self.calls.append(kwargs)
        return {
            "content": {
                "type": "text",
                "text": json.dumps(
                    {
                        "summary": "附件显示英伟达短期下跌与估值和供应链预期有关。",
                        "files": [
                            {
                                "file_id": "file-1",
                                "summary": "nvidia.pdf 讨论了股价压力和近期动作。",
                                "key_points": ["短期波动来自估值压力", "新产品和供应链仍是核心变量"],
                                "relevance": "可用于股票走势分析",
                                "relevance_score": 0.91,
                            }
                        ],
                    },
                    ensure_ascii=False,
                ),
            }
        }


class FakeOutlineSampling:
    def __init__(self, replies):
        self.replies = list(replies)
        self.calls = []

    def create_message(self, **kwargs):
        self.calls.append(kwargs)
        return {"content": {"type": "text", "text": self.replies.pop(0)}}


def outline_sampling_replies():
    return [
        json.dumps({"anchor_query": "NVIDIA recent decline outlook 2026", "facets": [{"task": "Explain the decline and assess the outlook"}]}),
        json.dumps({
            "queries": [
                {"text": "NVIDIA decline catalysts 2026", "covers": ["f1"]},
                {"text": "NVIDIA recent corporate actions 2026", "covers": ["f1"]},
                {"text": "NVIDIA valuation and outlook 2026", "covers": ["f1"]},
            ]
        }),
        json.dumps({
            "sections": [
                {"title": "Decline", "outline": "Explain recent catalysts.", "covers": ["f1"], "max_iterations": 5},
                {"title": "Actions", "outline": "Review company actions.", "covers": [], "max_iterations": 5},
                {"title": "Valuation", "outline": "Assess valuation and competition.", "covers": [], "max_iterations": 5},
                {"title": "Outlook", "outline": "Develop the forward scenario.", "covers": [], "max_iterations": 5},
            ]
        }),
    ]


# ---------------------------------------------------------------------------
# Settings & legacy migration
# ---------------------------------------------------------------------------


def test_mask_secret_uses_front_zero_back_four():
    assert mask_secret("tvly-abcdef1234") == "***1234"
    assert mask_secret("short") == "***hort"
    assert mask_secret("abc") == "***"
    assert mask_secret("") == ""


def test_settings_view(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    view = dispatcher.dispatch("app_get_settings", {})["settings"]
    assert view["tavily"]["configured"] is False
    assert view["research_root"] == str(tmp_path / ".research")


def test_legacy_tavily_key_migration_is_idempotent(tmp_path):
    root = tmp_path / ".research"
    settings = SettingsStore(root=root)
    settings.update(tavily_api_key="tvly-legacy-key-aaaa")
    creds = CredentialStore(root)

    assert migrate_legacy_tavily_key(settings, creds) is True
    assert "tavily_api_key" not in settings.read_raw()
    assert creds.get_token("tavily") == "tvly-legacy-key-aaaa"
    assert creds.status("tavily")["credential"] == "tvly-legacy-key-aaaa"

    # Re-running the migration is a no-op.
    assert migrate_legacy_tavily_key(settings, creds) is False
    assert creds.get_token("tavily") == "tvly-legacy-key-aaaa"


def test_credential_store_set_clear_remove(tmp_path):
    creds = CredentialStore(tmp_path / ".research")
    assert creds.status("tavily") == {"credential_status": "missing", "credential": ""}
    status = creds.set_token("tavily", "tvly-abcdef1234")
    assert status == {"credential_status": "configured", "credential": "tvly-abcdef1234"}
    creds.clear("tavily")
    assert creds.get_token("tavily") == ""
    creds.set_token("custom", "secret-zzzz")
    creds.remove("custom")
    assert creds.get_token("custom") == ""


# ---------------------------------------------------------------------------
# Job store v3
# ---------------------------------------------------------------------------


def test_job_create_update_latest_and_not_found(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    assert job["research_id"].startswith("research_")
    latest_status = dispatcher.dispatch("app_get_research_job", {})["job"]
    assert latest_status["research_id"] == job["research_id"]
    assert "job_transfer" not in latest_status
    payload = dispatcher.transfers.read(dispatcher.dispatch("app_get_research_job_payload", {"research_id": job["research_id"]})["transfer"])
    assert payload["job"]["research_id"] == job["research_id"]
    updated = dispatcher.dispatch(
        "app_update_research_job",
        {
            "research_id": job["research_id"],
            "updates": {
                "stage": "search_next_query",
                "progress": 25,
                "iteration": 1,
                "max_iterations": 5,
                "enabled_sources": ["tavily"],
            },
        },
    )
    assert updated["job"]["stage"] == "search_next_query"
    assert updated["job"]["progress"] == 25
    assert updated["job"]["iteration"] == 1
    assert updated["job"]["max_iterations"] == 5
    with pytest.raises(ValidationError):
        dispatcher.dispatch(
            "app_update_research_job",
            {"research_id": job["research_id"], "updates": {"tavily_api_key": "leak"}},
        )
    with pytest.raises(NotFoundError):
        dispatcher.dispatch("app_get_research_job", {"research_id": "missing"})


def test_job_store_writes_split_files_only_for_non_default_values(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="anna")
    research_id = job["research_id"]
    job_dir = jobs.job_dir_for(research_id)

    assert (job_dir / "job.json").exists()
    assert not (job_dir / "attachments.json").exists()
    assert not (job_dir / "attachment_context.json").exists()
    assert not (job_dir / "section_results.json").exists()
    assert jobs.load(research_id)["attachments"] == []

    jobs.update_metadata(research_id, {"attachments": [{"name": "brief.md"}]})

    assert (job_dir / "attachments.json").exists()
    assert jobs.load(research_id)["attachments"] == [{"name": "brief.md"}]

    jobs.update_metadata(research_id, {"attachments": []})

    assert not (job_dir / "attachments.json").exists()
    assert jobs.load(research_id)["attachments"] == []


def test_prepare_attachments_downloads_and_persists_text_chunks(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    result = dispatcher.dispatch(
        "app_prepare_attachments",
        {
            "research_id": job["research_id"],
            "attachments": [
                {
                    "name": "brief.md",
                    "path": "research-jobs/test/uploads/brief.md",
                    "content_type": "text/markdown",
                    "size_bytes": 38,
                    "download_url": "data:text/markdown;base64,QXR0YWNobWVudCByZXNlYXJjaCBicmllZiBmb3IgQW5uYS4=",
                }
            ],
        },
    )

    loaded = dispatcher.jobs.load(job["research_id"])
    context = loaded["attachment_context"]
    assert result["job"]["research_id"] == job["research_id"]
    assert result["job"]["attachment_context_summary"]["chunk_count"] == 1
    assert "attachment_context" not in result["job"]
    assert context["files"][0]["status"] == "ready"
    assert context["files"][0]["chunk_count"] == 1
    assert context["chunks"][0]["file_name"] == "brief.md"
    assert "Attachment research brief for Anna." in context["chunks"][0]["text"]


def test_prepare_image_attachment_uses_analysis_as_text_chunks(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    result = dispatcher.dispatch(
        "app_prepare_attachments",
        {
            "research_id": job["research_id"],
            "attachments": [
                {
                    "name": "chart.png",
                    "path": "research-jobs/test/uploads/chart.png",
                    "content_type": "image/png",
                    "size_bytes": 1024,
                    "download_url": "data:image/png;base64,iVBORw0KGgo=",
                    "image_analysis": {
                        "summary": "A line chart comparing quarterly revenue.",
                        "visible_text": "Revenue Q1 Q2",
                        "key_observations": ["Revenue rises across the chart."],
                        "chart_or_table": "Line chart",
                        "research_relevance": {"relevance": "Useful as visual evidence for the report.", "relevance_score": 0.8},
                        "uncertainties": ["Exact axis values are too small to read."],
                    },
                }
            ],
        },
    )

    loaded = dispatcher.jobs.load(job["research_id"])
    context = loaded["attachment_context"]
    assert result["job"]["attachment_context_summary"]["chunk_count"] == 0
    assert context["files"][0]["status"] == "ready"
    assert context["files"][0]["analysis"]["type"] == "image"
    assert context["files"][0]["chunk_count"] == 0
    assert context["files"][0]["local_path"].startswith("attachment-files/file-1-")
    assert (dispatcher.jobs.job_dir_for(job["research_id"]) / context["files"][0]["local_path"]).exists()
    assert context["files"][0]["analysis"]["summary"] == "A line chart comparing quarterly revenue."
    assert context["files"][0]["analysis"]["payload"]["visible_text"] == "Revenue Q1 Q2"

    selected = select_attachment_context(
        jobs=dispatcher.jobs,
        embeddings=FakeEmbeddings(),
        research_id=job["research_id"],
        query="quarterly revenue trend",
        top_k=4,
    )
    assert selected["selected_item_count"] == 1
    assert selected["selected_items"][0]["kind"] == "image_analysis"
    assert selected["selected_items"][0]["quote"] == ""
    assert "Image analysis JSON for chart.png" in selected["selected_context"]
    assert "file-1:image-summary" not in selected["selected_context"]
    assert '"visible_text": "Revenue Q1 Q2"' in selected["selected_context"]
    assert "Revenue rises across the chart." in selected["selected_context"]


def test_extract_docx_attachment_text_and_table(tmp_path):
    from docx import Document

    path = tmp_path / "brief.docx"
    document = Document()
    document.add_paragraph("英伟达股票研究附件")
    document.add_paragraph("数据中心收入增长，但估值压力上升。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "指标"
    table.rows[0].cells[1].text = "变化"
    document.save(path)

    text = extract_text(
        path,
        name="brief.docx",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    chunks = chunk_text(text)

    assert "英伟达股票研究附件" in text
    assert "指标 | 变化" in text
    assert chunks


def test_embed_attachment_chunks_batches_and_persists_vectors(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    context = {
        "version": 1,
        "prepared_at": "now",
        "files": [],
        "summary": "",
        "chunks": [
            {"chunk_id": f"file-1:{index + 1:04d}", "file_id": "file-1", "file_name": "brief.md", "index": index + 1, "text": f"chunk {index + 1}"}
            for index in range(18)
        ],
    }
    dispatcher.jobs.update_metadata(job["research_id"], {"attachment_context": context})
    fake = FakeEmbeddings()

    embed_attachment_chunks(jobs=dispatcher.jobs, embeddings=fake, research_id=job["research_id"])

    loaded = dispatcher.jobs.load(job["research_id"])["attachment_context"]
    assert [len(wave) for wave in fake.batch_calls] == [8, 8, 2]
    assert [len(call) for call in fake.calls] == [1] * 18
    assert loaded["embedding_status"] == "ready"
    assert all(chunk.get("embedding") for chunk in loaded["chunks"])


def test_embedding_client_honors_concurrency_limit_and_preserves_order():
    class TrackingEmbeddingsClient(AnnaEmbeddingsClient):
        def __init__(self):
            self.lock = threading.Lock()
            self.first_wave = threading.Barrier(MAX_PARALLEL_EMBEDDING_BATCHES)
            self.active = 0
            self.max_active = 0
            self.calls = []

        def create(self, *, texts, model="anna-managed-v1", timeout=30.0):
            batch_index = int(texts[0].removeprefix("batch-"))
            with self.lock:
                self.calls.append(batch_index)
                self.active += 1
                self.max_active = max(self.max_active, self.active)
            try:
                if batch_index < MAX_PARALLEL_EMBEDDING_BATCHES:
                    self.first_wave.wait(timeout=2)
                    time.sleep((MAX_PARALLEL_EMBEDDING_BATCHES - batch_index) * 0.002)
                return {"batch_index": batch_index, "data": [{"embedding": [float(batch_index)]}]}
            finally:
                with self.lock:
                    self.active -= 1

    client = TrackingEmbeddingsClient()
    batches = [[f"batch-{index}"] for index in range(MAX_PARALLEL_EMBEDDING_BATCHES + 1)]

    results = client.create_batches(batches=batches)

    assert client.max_active == MAX_PARALLEL_EMBEDDING_BATCHES
    assert [result["batch_index"] for result in results] == list(range(MAX_PARALLEL_EMBEDDING_BATCHES + 1))
    assert sorted(client.calls) == list(range(MAX_PARALLEL_EMBEDDING_BATCHES + 1))
    assert client.create_batches(batches=[]) == []


def test_embed_attachment_chunks_checkpoints_successes_and_retries_failures(tmp_path, monkeypatch):
    monkeypatch.setattr(attachment_embeddings_module, "EMBEDDING_RETRY_DELAY_SECONDS", 0)

    class ShortResultEmbeddings(FakeEmbeddings):
        def __init__(self):
            super().__init__()
            self.rounds = 0

        def create_batches_settled(self, *, batches, model="anna-managed-v1", timeout=30.0):
            self.rounds += 1
            return [
                EmbeddingBatchOutcome(result={
                    "data": [{"embedding": [1.0, float(index + 1)]} for index, _ in enumerate(batch)],
                    "_meta": {"dimensions": 2},
                })
                if self.rounds > 1 or "chunk 3" not in batch
                else EmbeddingBatchOutcome(result={"data": [], "_meta": {"dimensions": 2}})
                for batch in batches
            ]

    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    context = {
        "version": 1,
        "prepared_at": "now",
        "files": [],
        "summary": "",
        "chunks": [
            {"chunk_id": f"file-1:{index + 1:04d}", "file_id": "file-1", "file_name": "brief.md", "index": index + 1, "text": f"chunk {index + 1}"}
            for index in range(3)
        ],
    }
    dispatcher.jobs.update_metadata(job["research_id"], {"attachment_context": context})

    transient = ShortResultEmbeddings()
    embed_attachment_chunks(jobs=dispatcher.jobs, embeddings=transient, research_id=job["research_id"])

    loaded = dispatcher.jobs.load(job["research_id"])["attachment_context"]
    assert transient.rounds == 2
    assert loaded["embedding_status"] == "ready"
    assert all(chunk.get("embedding") for chunk in loaded["chunks"])

    persistent_job = dispatcher.dispatch("app_create_research_job", {"query": "persistent embedding failure"})["job"]
    persistent_context = {
        "version": 1,
        "prepared_at": "now",
        "files": [],
        "summary": "",
        "chunks": [
            {
                "chunk_id": f"file-2:{index + 1:04d}",
                "file_id": "file-2",
                "file_name": "persistent.md",
                "index": index + 1,
                "text": f"chunk {index + 1}",
            }
            for index in range(3)
        ],
    }
    dispatcher.jobs.update_metadata(persistent_job["research_id"], {"attachment_context": persistent_context})

    class PersistentFailureEmbeddings(FakeEmbeddings):
        def create_batches_settled(self, *, batches, model="anna-managed-v1", timeout=30.0):
            return [
                EmbeddingBatchOutcome(result={"data": []})
                if "chunk 3" in batch
                else EmbeddingBatchOutcome(result=self.create(texts=batch, model=model, timeout=timeout))
                for batch in batches
            ]

    with pytest.raises(EmbeddingsError, match="still failed after automatic retry"):
        embed_attachment_chunks(jobs=dispatcher.jobs, embeddings=PersistentFailureEmbeddings(), research_id=persistent_job["research_id"])
    partial = dispatcher.jobs.load(persistent_job["research_id"])["attachment_context"]
    assert all(chunk.get("embedding") for chunk in partial["chunks"][:2])
    assert not partial["chunks"][2].get("embedding")
    assert partial["embedding_status"] == "partial"


def test_embed_attachment_chunks_skips_existing_vectors(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    context = {
        "version": 1,
        "prepared_at": "now",
        "files": [],
        "summary": "",
        "chunks": [
            {
                "chunk_id": f"file-1:{index + 1:04d}",
                "file_id": "file-1",
                "file_name": "brief.md",
                "index": index + 1,
                "text": f"chunk {index + 1}",
                **({"embedding": [9.0, 9.0]} if index == 0 else {}),
            }
            for index in range(4)
        ],
    }
    dispatcher.jobs.update_metadata(job["research_id"], {"attachment_context": context})
    fake = FakeEmbeddings()

    embed_attachment_chunks(jobs=dispatcher.jobs, embeddings=fake, research_id=job["research_id"])

    loaded = dispatcher.jobs.load(job["research_id"])["attachment_context"]
    assert fake.calls == [["chunk 2", "chunk 3"], ["chunk 4"]]
    assert loaded["chunks"][0]["embedding"] == [9.0, 9.0]
    assert all(chunk.get("embedding") for chunk in loaded["chunks"])


def test_summarize_attachment_context_selects_top_chunks_and_writes_per_file_summary(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "nvidia stock"})["job"]
    context = {
        "version": 1,
        "prepared_at": "now",
        "files": [{"id": "file-1", "name": "nvidia.pdf", "status": "ready", "chunk_count": 2}],
        "summary": "legacy summary",
        "chunks": [
            {"chunk_id": "file-1:0001", "file_id": "file-1", "file_name": "nvidia.pdf", "index": 1, "text": "Nvidia stock fell on valuation concerns.", "embedding": [1.0, 1.0]},
            {"chunk_id": "file-1:0002", "file_id": "file-1", "file_name": "nvidia.pdf", "index": 2, "text": "Supply chain and product launch details.", "embedding": [0.1, 0.1]},
        ],
    }
    dispatcher.jobs.update_metadata(job["research_id"], {"attachment_context": context})
    sampling = FakeSampling()

    summarize_attachment_context(
        jobs=dispatcher.jobs,
        embeddings=FakeEmbeddings(),
        sampling=sampling,
        research_id=job["research_id"],
        query="why did nvidia stock fall",
        top_k=1,
        invoke_id="invoke-test",
    )

    loaded = dispatcher.jobs.load(job["research_id"])["attachment_context"]
    assert loaded["summary_status"] == "ready"
    assert loaded["summary_mode"] == "ai_topk_by_file"
    assert loaded["files"][0]["analysis"]["summary"] == "nvidia.pdf 讨论了股价压力和近期动作。"
    assert loaded["files"][0]["analysis"]["relevance_score"] == 0.91
    assert loaded["files"][0]["analysis"]["selected_chunk_ids"] == ["file-1:0001"]
    assert "附件显示英伟达" in loaded["summary"]
    prompt = sampling.calls[0]["messages"][0]["content"]["text"]
    assert "Nvidia stock fell" in prompt
    assert '"relevance_score":"number from 0 to 1"' in prompt
    assert "Set relevance_score from 0 to 1" in prompt
    assert "0.0-0.2 = unrelated" in prompt
    assert "0.76-1.0 = directly relevant evidence" in prompt


def test_select_attachment_context_returns_ranked_chunk_text(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "nvidia stock"})["job"]
    context = {
        "version": 1,
        "prepared_at": "now",
        "files": [{"id": "file-1", "name": "nvidia.pdf", "status": "ready", "chunk_count": 1, "analysis": {"type": "text", "relevance_score": 0.8}}],
        "summary": "AI summary should not be the writing context",
        "chunks": [
            {"chunk_id": "file-1:0001", "file_id": "file-1", "file_name": "nvidia.pdf", "index": 1, "text": "Original filing chunk about Nvidia supply constraints.", "embedding": [1.0, 1.0]},
        ],
    }
    dispatcher.jobs.update_metadata(job["research_id"], {"attachment_context": context})

    selected = select_attachment_context(
        jobs=dispatcher.jobs,
        embeddings=FakeEmbeddings(),
        research_id=job["research_id"],
        query="nvidia supply constraints",
        top_k=1,
    )

    assert selected["selected_item_count"] == 1
    assert selected["selected_items"][0]["kind"] == "chunk"
    assert selected["selected_items"][0]["quote"] == "Original filing chunk about Nvidia supply constraints."
    assert "Use this analysis as supporting evidence only when the claim is directly grounded in visible content." in selected["selected_context"]
    assert "Original filing chunk" in selected["selected_context"]
    assert "file-1:0001" not in selected["selected_context"]
    assert "AI summary should not" not in selected["selected_context"]


def test_select_attachment_context_skips_irrelevant_files(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "us gdp forecast"})["job"]
    context = {
        "version": 1,
        "prepared_at": "now",
        "files": [
            {"id": "file-1", "name": "macro.pdf", "status": "ready", "chunk_count": 1, "analysis": {"type": "text", "relevance_score": 0.8}},
            {"id": "file-2", "name": "gpu.pdf", "status": "ready", "chunk_count": 1, "analysis": {"type": "text", "relevance_score": 0.05, "relevance": "无相关性。"}},
        ],
        "summary": "attachment summary",
        "chunks": [
            {"chunk_id": "file-1:0001", "file_id": "file-1", "file_name": "macro.pdf", "index": 1, "text": "US GDP forecast and labor market data.", "embedding": [1.0, 1.0]},
            {"chunk_id": "file-2:0001", "file_id": "file-2", "file_name": "gpu.pdf", "index": 1, "text": "GPU supply chain and semiconductor demand.", "embedding": [1.0, 1.0]},
        ],
    }
    dispatcher.jobs.update_metadata(job["research_id"], {"attachment_context": context})

    selected = select_attachment_context(
        jobs=dispatcher.jobs,
        embeddings=FakeEmbeddings(),
        research_id=job["research_id"],
        query="us gdp forecast",
        top_k=4,
    )

    assert selected["selected_item_count"] == 1
    assert "US GDP forecast" in selected["selected_context"]
    assert "GPU supply chain" not in selected["selected_context"]


def test_get_research_job_requires_aps_transfer_capability(tmp_path):
    root = tmp_path / ".research"
    dispatcher = AppDispatcher(
        settings=SettingsStore(root=root),
        jobs=JobStore(root=root),
        selector=LexicalContextSelector(max_sources=4, context_budget=4000),
    )
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    with pytest.raises(ValidationError, match="APS Files transfer capability is unavailable"):
        dispatcher.dispatch("app_get_research_job", {"research_id": job["research_id"]})


def test_outline_discovery_runs_sampling_and_search_entirely_in_backend(tmp_path, monkeypatch):
    monkeypatch.setenv("ANNA_RESEARCHER_FAKE_TAVILY", "1")
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "NVIDIA recent decline, company actions, and outlook"})["job"]
    research_id = job["research_id"]
    dispatcher.dispatch(
        "app_save_confirmed_research_role",
        {"research_id": research_id, "role": {"server": "Analyst", "agent_role_prompt": "Use current market evidence."}},
    )
    dispatcher.jobs.update_metadata(research_id, {"attachment_context": {
        "summary": "Uploaded earnings memo",
        "files": [{
            "id": "file-1",
            "name": "earnings-memo.pdf",
            "status": "ready",
            "analysis": {
                "summary": "The memo says Blackwell shipment timing remains uncertain.",
                "key_points": ["Shipment timing needs independent confirmation."],
                "relevance": "Directly relevant to recent company actions.",
                "relevance_score": 0.9,
                "payload": {"private_detail": "must not enter search planning"},
            },
        }],
    }})
    sampling = FakeOutlineSampling(outline_sampling_replies())

    result = generate_outline_draft(
        dispatcher=dispatcher,
        sampling=sampling,
        research_id=research_id,
        source_ids=["tavily"],
        invoke_id="invoke-outline",
    )

    assert len(result["outline"]) == 4
    assert "selected_context" not in result
    assert "context_transfer" not in result
    assert len(sampling.calls) == 3
    assert "Blackwell shipment timing remains uncertain" in sampling.calls[0]["messages"][0]["content"]["text"]
    assert "prioritize missing context, independent corroboration" in sampling.calls[1]["messages"][0]["content"]["text"]
    assert "private_detail" not in json.dumps(sampling.calls)
    assert "Seed search context" in sampling.calls[1]["messages"][0]["content"]["text"]
    assert "Selected web discovery context" in sampling.calls[2]["messages"][0]["content"]["text"]
    assert all(call["metadata"]["executa_invoke_id"] == "invoke-outline" for call in sampling.calls)
    discovery = dispatcher.jobs.load(research_id)["outline_discovery"]
    assert discovery["status"] == "completed"
    assert discovery["seed"]["raw_results"]
    assert discovery["research_calls"][0]["raw_results"]
    assert discovery["selected_contexts"]["research"]["selected_sources"]


def test_outline_regeneration_reuses_persisted_discovery_without_http_context(tmp_path, monkeypatch):
    monkeypatch.setenv("ANNA_RESEARCHER_FAKE_TAVILY", "1")
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "NVIDIA recent decline and outlook"})["job"]
    research_id = job["research_id"]
    dispatcher.dispatch(
        "app_save_confirmed_research_role",
        {"research_id": research_id, "role": {"server": "Analyst", "agent_role_prompt": "Use current market evidence."}},
    )
    generate_outline_draft(
        dispatcher=dispatcher,
        sampling=FakeOutlineSampling(outline_sampling_replies()),
        research_id=research_id,
        source_ids=["tavily"],
    )
    before = dispatcher.jobs.load(research_id)["outline_discovery"]
    outline_reply = outline_sampling_replies()[-1]
    sampling = FakeOutlineSampling([outline_reply])

    result = generate_outline_draft(
        dispatcher=dispatcher,
        sampling=sampling,
        research_id=research_id,
        source_ids=["tavily"],
        instruction="Emphasize scenario boundaries.",
        reuse_discovery=True,
    )

    assert len(result["outline"]) == 4
    assert len(sampling.calls) == 1
    assert "Regeneration requirement: Emphasize scenario boundaries." in sampling.calls[0]["messages"][0]["content"]["text"]
    after = dispatcher.jobs.load(research_id)["outline_discovery"]
    assert len(after["research_calls"]) == len(before["research_calls"])


def test_compact_job_view_exposes_v3_fields(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    immediate = dispatcher.dispatch("app_get_research_job", {"research_id": job["research_id"]})["job"]
    loaded = immediate
    assert loaded["schema_version"] == 5
    assert loaded["iterations"] == []
    assert loaded["research_log"] == []
    assert loaded["iteration"] == 0
    assert loaded["max_iterations"] == 5
    assert loaded["enabled_sources"] == []


def test_compact_job_view_omits_full_section_citation_sources(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="anna")
    research_id = job["research_id"]
    jobs.save_section_result(
        research_id,
        "section-1",
        {
            "status": "completed",
            "section_markdown": "## Section\n\nEvidence [1][2]",
            "section_summary": "Summary",
            "subsection_headers": ["Evidence basis", "Risk analysis"],
            "source_urls": ["https://example.test/a"],
            "citation_sources": [
                {"kind": "url", "url": "https://example.test/a", "title": "A", "content": "x" * 1000},
                {"kind": "attachment", "file_id": "file-1", "file_name": "a.pdf", "chunk_id": "file-1:0001", "quote": "y" * 1000},
            ],
        },
    )

    view = compact_job_view(jobs.load(research_id))
    section = view["section_results"]["section-1"]

    assert "section_markdown" not in section
    assert "source_urls" not in section
    assert "citation_sources" not in section
    assert section["section_markdown_chars"] == len("## Section\n\nEvidence [1][2]")
    assert section["subsection_headers"] == ["Evidence basis", "Risk analysis"]
    assert section["source_count"] == 1
    assert section["citation_source_count"] == 2
    assert section["attachment_citation_count"] == 1


def test_job_store_has_called_dedup_uses_normalized_query(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="anna")
    research_id = job["research_id"]
    assert jobs.has_called(research_id, "tavily", normalize_query_for_dedup("Anna  App")) is False
    jobs.append_iteration(
        research_id,
        iteration=1,
        source_id="tavily",
        source_name="Tavily",
        queries=["Anna App"],
        source_calls=[{"query": "Anna App", "items": [{"url": "https://x.example", "title": "t"}], "duration_ms": 1}],
        raw_results=[{"query": "Anna App", "url": "https://x.example", "title": "t", "content": "c"}],
    )
    loaded = jobs.load(research_id)
    assert "items" not in loaded["iterations"][0]["source_calls"][0]
    assert loaded["iterations"][0]["source_calls"][0]["results_count"] == 1
    assert loaded["iterations"][0]["source_calls"][0]["top_titles"] == ["t"]
    assert loaded["iterations"][0]["raw_results"][0]["url"] == "https://x.example"
    assert loaded["iterations"][0]["raw_results"][0]["content"] == "c"
    assert jobs.has_called(research_id, "tavily", normalize_query_for_dedup("anna  app")) is True
    assert jobs.has_called(research_id, "tavily", normalize_query_for_dedup("different")) is False


def test_normalize_query_collapses_whitespace_and_case():
    assert normalize_query_for_dedup("  Anna   App  ") == "anna app"
    assert normalize_query_for_dedup("ANNA\tApp") == "anna app"
    assert normalize_query_for_dedup("") == ""


# ---------------------------------------------------------------------------
# Native DuckDuckGo adapter
# ---------------------------------------------------------------------------


def test_duckduckgo_native_search_normalizes_results(monkeypatch):
    class FakeDDGS:
        def text(self, query, *, region, max_results):
            assert query == "anna research"
            assert region == "us-en"
            assert max_results == 20
            return [
                {"href": "https://example.com/a", "title": "Alpha", "body": "alpha body"},
                {"url": "https://example.com/b", "title": "Beta", "snippet": "beta snippet"},
                {"href": "https://example.com/c", "content": "content fallback"},
                {"title": "No useful content"},
            ]

    monkeypatch.setattr(duckduckgo_native, "_create_client", lambda: FakeDDGS())

    results = duckduckgo_native.search_duckduckgo(" anna research ", max_results=50, region="us-en")

    assert results == [
        {
            "query": "anna research",
            "url": "https://example.com/a",
            "title": "Alpha",
            "content": "alpha body",
        },
        {
            "query": "anna research",
            "url": "https://example.com/b",
            "title": "Beta",
            "content": "beta snippet",
        },
        {
            "query": "anna research",
            "url": "https://example.com/c",
            "title": "https://example.com/c",
            "content": "content fallback",
        },
    ]


def test_duckduckgo_native_search_handles_empty_query(monkeypatch):
    monkeypatch.setattr(duckduckgo_native, "_create_client", lambda: pytest.fail("client should not be created"))

    assert duckduckgo_native.search_duckduckgo("   ") == []


def test_duckduckgo_native_search_wraps_ddgs_errors(monkeypatch):
    class FailingDDGS:
        def text(self, query, *, region, max_results):
            raise RuntimeError("blocked")

    monkeypatch.setattr(duckduckgo_native, "_create_client", lambda: FailingDDGS())

    with pytest.raises(duckduckgo_native.DuckDuckGoSearchError, match="duckduckgo search failed"):
        duckduckgo_native.search_duckduckgo("anna")


def test_native_research_source_executor_call_wraps_ddgs_results():
    executor = NativeResearchSourceExecutor(
        clock=_fixed_clock(),
        adapters={
            "ddgs": lambda query: [
                {"query": query, "url": "https://example.com/a", "title": "Alpha", "content": "alpha body"},
            ]
        },
        extractor=lambda items, **kwargs: items,
    )

    result = executor.call(_native_definition(), " anna ")

    assert result.source_id == "duckduckresearch"
    assert result.source_name == "DuckDuckResearch"
    assert result.query == "anna"
    assert result.duration_ms == 125
    assert result.error is None
    assert result.items == [
        {
            "query": "anna",
            "url": "https://example.com/a",
            "title": "Alpha",
            "content": "alpha body",
            "source_id": "duckduckresearch",
            "source_name": "DuckDuckResearch",
        }
    ]


def test_native_research_source_executor_extracts_by_default():
    extractor_calls = []

    def fake_extractor(items, **kwargs):
        extractor_calls.append((items, kwargs))
        enriched = []
        for item in items:
            enriched.append({**item, "raw_content": "Fetched body", "content": item["content"] + "\n\nFull content:\nFetched body"})
        return enriched

    executor = NativeResearchSourceExecutor(
        clock=_fixed_clock(),
        adapters={
            "ddgs": lambda query: [
                {"query": query, "url": "https://example.com/a", "title": "Alpha", "content": "alpha snippet"},
            ]
        },
        extractor=fake_extractor,
    )
    definition = _native_definition()
    definition["native"]["max_urls"] = 2
    definition["native"]["max_pdf_pages"] = 4
    definition["native"].pop("browser_fallback", None)
    definition["native"]["browser_fallback_min_chars"] = 250
    definition["native"]["browser_timeout"] = 12

    result = executor.call(definition, "anna")

    assert result.error is None
    assert result.items[0]["raw_content"] == "Fetched body"
    assert result.items[0]["content"] == "alpha snippet\n\nFull content:\nFetched body"
    assert extractor_calls == [
        (
            [{"query": "anna", "url": "https://example.com/a", "title": "Alpha", "content": "alpha snippet"}],
            {
                "query": "anna",
                "max_urls": 2,
                "timeout": 20.0,
                "max_pdf_pages": 4,
                "browser_fallback": True,
                "browser_fallback_min_chars": 250,
                "browser_timeout": 12.0,
                "page_cache": None,
            },
        )
    ]


def test_native_research_source_executor_reports_extraction_failure():
    def failing_extractor(items, **kwargs):
        raise RuntimeError("fetch blocked")

    executor = NativeResearchSourceExecutor(
        clock=_fixed_clock(),
        adapters={"ddgs": lambda query: [{"query": query, "url": "https://example.com/a", "title": "Alpha", "content": "alpha"}]},
        extractor=failing_extractor,
    )
    definition = _native_definition()

    result = executor.call(definition, "anna")
    test = executor.test(definition, "anna")

    assert result.error == "upstream_5xx"
    assert result.items == []
    assert test.error and test.error["code"] == "upstream_5xx"
    assert "native extraction failed" in test.error["message"]


def test_native_research_source_executor_reports_empty_results():
    executor = NativeResearchSourceExecutor(clock=_fixed_clock(), adapters={"ddgs": lambda query: []})

    result = executor.call(_native_definition(), "anna")
    test = executor.test(_native_definition(), "anna")

    assert result.error == "empty_result"
    assert result.items == []
    assert test.error == {"code": "empty_result", "message": "native source returned no results"}
    assert test.extracted == []


def test_native_research_source_executor_reports_bad_definition():
    executor = NativeResearchSourceExecutor(clock=_fixed_clock())

    result = executor.call({"id": "native", "name": "Native", "native": {"adapter": "missing"}}, "anna")
    test = executor.test({"id": "native", "name": "Native", "native": {"adapter": "missing"}}, "anna")

    assert result.error == "bad_definition"
    assert result.items == []
    assert test.error and test.error["code"] == "bad_definition"


def test_pdf_extraction_reads_local_pdf(tmp_path):
    fitz = pytest.importorskip("fitz")
    pdf_path = tmp_path / "sample.pdf"
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Anna PDF extraction smoke test")
    doc.save(str(pdf_path))
    doc.close()

    result = extract_pdf(str(pdf_path), max_pages=1, max_chars_per_page=2000)

    assert result.status == "success"
    assert result.content_type == "pdf"
    assert result.url == str(pdf_path)
    assert "Anna PDF extraction smoke test" in result.raw_content


def test_pdf_extraction_does_not_truncate_pages_at_legacy_character_limit(monkeypatch):
    page_text = "complete pdf text " * 700

    class FakePage:
        def get_text(self, mode):
            assert mode == "text"
            return page_text

    class FakeDocument:
        metadata = {"title": "Full PDF"}

        def __len__(self):
            return 1

        def load_page(self, index):
            assert index == 0
            return FakePage()

        def close(self):
            pass

    monkeypatch.setattr("fitz.open", lambda path: FakeDocument())

    title, content = pdf_extraction._extract_pdf_file("full.pdf", max_pages=None, max_chars_per_page=8000)

    assert title == "Full PDF"
    assert len(content) > 8000
    assert content == page_text.strip()


def test_arxiv_extraction_reads_paper_metadata(monkeypatch):
    def fake_load(query, *, max_results):
        assert query == "2401.12345"
        assert max_results == 2
        return {
            "title": "Anna Research Systems",
            "authors": "Ada Lovelace, Alan Turing",
            "published": "2026-06-22T00:00:00+00:00",
            "summary": "A paper about research agents.",
        }

    monkeypatch.setattr(arxiv_extraction, "_load_arxiv_document", fake_load)

    result = arxiv_extraction.extract_arxiv("https://arxiv.org/abs/2401.12345v2")

    assert result.status == "success"
    assert result.content_type == "arxiv"
    assert result.title == "Anna Research Systems"
    assert "Published: 2026-06-22T00:00:00+00:00" in result.raw_content
    assert "Author: Ada Lovelace, Alan Turing" in result.raw_content
    assert "Content: A paper about research agents." in result.raw_content


def test_arxiv_extraction_reports_empty_source():
    result = arxiv_extraction.extract_arxiv(" ")

    assert result.status == "failed"
    assert result.error == "empty_source"


def test_html_extraction_reads_local_html_and_removes_noise(tmp_path):
    html_path = tmp_path / "page.html"
    html_path.write_text(
        """
        <html>
          <head><title>Anna Page</title><script>bad()</script></head>
          <body>
            <nav>Navigation should disappear</nav>
            <article>
              <h1>Research Findings</h1>
              <p>Anna extracts the useful article body.</p>
              <p>Read the <a href="https://example.com/source">source page</a> for details.</p>
              <ul><li>First finding</li><li>Second finding</li></ul>
            </article>
            <footer>Footer should disappear</footer>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    result = extract_html(str(html_path), max_chars_per_page=2000)

    assert result.status == "success"
    assert result.content_type == "html"
    assert result.title == "Anna Page"
    assert "# Research Findings" in result.raw_content
    assert "Anna extracts the useful article body." in result.raw_content
    assert "source page⟨1⟩" in result.raw_content
    assert "## References" in result.raw_content
    assert "⟨1⟩ https://example.com/source: source page" in result.raw_content
    assert "- First finding" in result.raw_content
    assert "Navigation should disappear" not in result.raw_content
    assert "Footer should disappear" not in result.raw_content


def test_html_extraction_prunes_but_does_not_filter_or_truncate_by_query(tmp_path):
    html_path = tmp_path / "filtered.html"
    html_path.write_text(
        """
        <html>
          <head><title>Market Page</title></head>
          <body>
            <main>
              <section>
                <h2>Anthropic financing</h2>
                <p>Anthropic raised new funding from strategic cloud partners and expanded its enterprise market position.</p>
                <p>Claude adoption in enterprise accounts supports the company's market positioning.</p>
              </section>
              <section>
                <h2>Recipe notes</h2>
                <p>Banana bread tastes better with cinnamon and walnuts.</p>
              </section>
              <p><a href="/a">Home</a> <a href="/b">Login</a> <a href="/c">Register</a> <a href="/d">Share</a></p>
            </main>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    result = extract_html(
        str(html_path),
        max_chars_per_page=4000,
        query="Anthropic financing enterprise market positioning",
    )

    assert result.status == "success"
    assert "# Anthropic financing" in result.raw_content
    assert "strategic cloud partners" in result.raw_content
    assert "enterprise accounts" in result.raw_content
    assert "Banana bread" in result.raw_content
    assert "Home" not in result.raw_content


def test_html_extraction_preserves_content_beyond_configured_legacy_limit(tmp_path):
    html_path = tmp_path / "long.html"
    body = "complete extraction marker " * 500
    html_path.write_text(f"<html><body><main><p>{body}</p></main></body></html>", encoding="utf-8")

    result = extract_html(str(html_path), max_chars_per_page=8000, query="unrelated query")

    assert result.status == "success"
    assert len(result.raw_content) > 8000
    assert result.raw_content.endswith("marker")


def test_fetch_url_uses_browser_fallback_for_short_static_content(monkeypatch):
    calls = []

    monkeypatch.setattr(
        extraction_fetcher,
        "extract_html",
        lambda url, **kwargs: calls.append(("static", url, kwargs)) or ExtractedPage(url=url, title="Short", raw_content="Short", content_type="html"),
    )
    monkeypatch.setattr(
        extraction_fetcher,
        "extract_with_browser_fallback",
        lambda url, **kwargs: calls.append(("browser", url, kwargs))
        or ExtractedPage(url=url, title="Browser", raw_content="# Browser markdown\n\nThis dynamic page contains enough useful body text.", content_type="html"),
    )

    page = extraction_fetcher.fetch_url(
        "https://example.com/dynamic",
        query="dynamic page",
        browser_fallback=True,
        browser_fallback_min_chars=20,
        browser_timeout=9,
    )

    assert page.title == "Browser"
    assert page.raw_content == "# Browser markdown\n\nThis dynamic page contains enough useful body text."
    assert calls == [
        ("static", "https://example.com/dynamic", {"timeout": 20.0, "max_chars_per_page": 12000, "query": "dynamic page"}),
        ("browser", "https://example.com/dynamic", {"query": "dynamic page", "timeout": 9, "max_chars_per_page": 12000}),
    ]


def test_browser_fallback_prefers_fit_markdown_then_raw_markdown():
    class Markdown:
        fit_markdown = "  "
        raw_markdown = "\n# Full page markdown\n\nUseful body.\n"

        def __str__(self):
            return "object fallback should not be used"

    assert browser_fallback._markdown_text(Markdown()) == "# Full page markdown\n\nUseful body."
    Markdown.fit_markdown = "\n# Filtered markdown\n"
    assert browser_fallback._markdown_text(Markdown()) == "# Full page markdown\n\nUseful body."


def test_browser_fallback_does_not_truncate_complete_markdown(monkeypatch):
    markdown = "complete browser markdown " * 500

    async def fake_extract(url, *, query, timeout):
        return "Full browser page", "", markdown

    monkeypatch.setattr(browser_fallback, "_extract_with_crawl4ai", fake_extract)

    result = browser_fallback.extract_with_browser_fallback("https://example.com/full", max_chars_per_page=100)

    assert result.status == "success"
    assert result.raw_content == markdown.strip()

def test_browser_fallback_uses_cleaned_html_when_markdown_is_empty():
    class Crawler:
        async def arun(self, *, url, config):
            return type(
                "Result",
                (),
                {
                    "success": True,
                    "markdown": None,
                    "cleaned_html": "\n<main>Cleaned page body</main>\n",
                    "metadata": {"title": "Cleaned", "icon": "https://example.com/icon.png"},
                },
            )()

    title, icon, markdown = asyncio.run(browser_fallback._crawl_one(Crawler(), "https://example.com", config=object()))

    assert title == "Cleaned"
    assert icon == "https://example.com/icon.png"
    assert markdown == "<main>Cleaned page body</main>"


def test_tavily_uses_long_summary_and_crawls_only_short_content():
    crawl_calls = []

    def browser_extractor(urls, **kwargs):
        crawl_calls.append((urls, kwargs))
        return [
            ExtractedPage(
                url=url,
                title="Crawled title",
                icon="https://short.example/favicon.ico",
                raw_content="# Crawled full page\n\n" + ("Useful body with research evidence. " * 5),
                content_type="html",
            )
            for url in urls
        ]

    long_summary = "L" * (TAVILY_PREFETCH_MIN_CHARS + 1)
    exact_threshold = "S" * TAVILY_PREFETCH_MIN_CHARS
    enriched = enrich_tavily_items(
        [
            {"url": "https://long.example/a", "title": "Long", "content": long_summary},
            {"url": "https://short.example/b", "title": "Short", "content": exact_threshold},
        ],
        browser_extractor=browser_extractor,
    )

    assert enriched[0]["raw_content"] == long_summary
    assert enriched[0]["content_type"] == "tavily_summary"
    assert enriched[0]["extraction_status"] == "success"
    assert enriched[1]["raw_content"].startswith("# Crawled full page")
    assert enriched[1]["content_type"] == "html"
    assert enriched[1]["title"] == "Short"
    assert enriched[1]["icon"] == "https://short.example/favicon.ico"
    assert crawl_calls == [(["https://short.example/b"], {"timeout": 15.0})]


def test_tavily_falls_back_to_short_summary_when_browser_extraction_fails():
    url = "https://short.example/fallback"
    cache = {
        same_url_without_fragment(url): ExtractedPage(
            url=url,
            content_type="html",
            status="failed",
            error="old failure",
        )
    }
    crawl_calls = []

    def browser_extractor(urls, **kwargs):
        crawl_calls.append((urls, kwargs))
        return [ExtractedPage(url=item, content_type="html", status="failed", error="timeout") for item in urls]

    enriched = enrich_tavily_items(
        [{"url": url, "title": "Short result", "content": "A useful but short Tavily summary."}],
        page_cache=cache,
        browser_extractor=browser_extractor,
    )

    assert crawl_calls == [([url], {"timeout": 15.0})]
    assert enriched[0]["extraction_status"] == "success"
    assert enriched[0]["content_type"] == "tavily_summary_fallback"
    assert enriched[0]["raw_content"] == "A useful but short Tavily summary."
    assert enriched[0]["extraction_fallback_reason"] == "timeout"
    assert cache[same_url_without_fragment(url)].content_type == "tavily_summary_fallback"


def test_tavily_does_not_cache_failed_browser_result_without_summary():
    url = "https://empty.example/article"
    cache = {}

    enriched = enrich_tavily_items(
        [{"url": url, "title": "Empty", "content": ""}],
        page_cache=cache,
        browser_extractor=lambda urls, **kwargs: [
            ExtractedPage(url=item, content_type="html", status="failed", error="timeout") for item in urls
        ],
    )

    assert enriched[0]["extraction_status"] == "failed"
    assert enriched[0]["extraction_error"] == "timeout"
    assert same_url_without_fragment(url) not in cache


def test_tavily_reuses_existing_web_document_before_summary_threshold():
    url = "https://cached.example/article#section"
    cache = {
        same_url_without_fragment(url): ExtractedPage(
            url="https://cached.example/article",
            title="Cached",
            raw_content="Previously crawled full page",
            content_type="html",
        )
    }

    enriched = enrich_tavily_items(
        [{"url": url, "title": "Tavily title", "content": "short"}],
        page_cache=cache,
        browser_extractor=lambda *_args, **_kwargs: pytest.fail("cached URL should not be crawled"),
    )

    assert enriched[0]["raw_content"] == "Previously crawled full page"
    assert enriched[0]["content_type"] == "html"


def test_fetch_url_marks_low_value_fallback_result_as_failed(monkeypatch):
    monkeypatch.setattr(
        extraction_fetcher,
        "extract_html",
        lambda url, **kwargs: ExtractedPage(url=url, title="Shell", raw_content="html", content_type="html"),
    )
    monkeypatch.setattr(
        extraction_fetcher,
        "extract_with_browser_fallback",
        lambda url, **kwargs: ExtractedPage(url=url, title="Shell", raw_content="Please wait while your request is being verified...", content_type="html"),
    )

    page = extraction_fetcher.fetch_url(
        "https://example.com/shell",
        browser_fallback=True,
        browser_fallback_min_chars=20,
    )

    assert page.status == "failed"
    assert page.error == "low_value_content"


def test_fetch_many_batches_browser_fallback(monkeypatch):
    calls = []

    monkeypatch.setattr(
        extraction_fetcher,
        "extract_html",
        lambda url, **kwargs: calls.append(("static", url, kwargs)) or ExtractedPage(url=url, title="Short", raw_content="Short", content_type="html"),
    )
    monkeypatch.setattr(
        extraction_fetcher,
        "extract_many_with_browser_fallback",
        lambda urls, **kwargs: calls.append(("browser_many", list(urls), kwargs))
        or [
            ExtractedPage(url=url, title="Browser", raw_content=f"# Browser markdown for {url}\n\nUseful dynamic body text.", content_type="html")
            for url in urls
        ],
    )

    pages = extraction_fetcher.fetch_many(
        ["https://example.com/a", "https://example.com/b"],
        query="dynamic page",
        browser_fallback=True,
        browser_fallback_min_chars=20,
        browser_timeout=9,
    )

    assert [page.url for page in pages] == ["https://example.com/a", "https://example.com/b"]
    assert all(page.status == "success" for page in pages)
    assert calls == [
        ("static", "https://example.com/a", {"timeout": 20.0, "max_chars_per_page": 12000, "query": "dynamic page"}),
        ("static", "https://example.com/b", {"timeout": 20.0, "max_chars_per_page": 12000, "query": "dynamic page"}),
        (
            "browser_many",
            ["https://example.com/a", "https://example.com/b"],
            {"query": "dynamic page", "timeout": 9, "max_chars_per_page": 12000},
        ),
    ]


def test_fetch_many_uses_page_cache(monkeypatch):
    calls = []

    monkeypatch.setattr(
        extraction_fetcher,
        "extract_html",
        lambda url, **kwargs: calls.append(("static", url)) or ExtractedPage(url=url, title="Cached", raw_content="Long enough cached body text.", content_type="html"),
    )

    cache = {}
    first = extraction_fetcher.fetch_many(
        ["https://example.com/a"],
        browser_fallback=True,
        browser_fallback_min_chars=20,
        page_cache=cache,
    )
    second = extraction_fetcher.fetch_many(
        ["https://example.com/a#again"],
        browser_fallback=True,
        browser_fallback_min_chars=20,
        page_cache=cache,
    )

    assert first[0] is second[0]
    assert calls == [("static", "https://example.com/a")]


def test_extraction_fetcher_routes_by_url_type(monkeypatch):
    calls = []

    monkeypatch.setattr(
        extraction_fetcher,
        "extract_arxiv",
        lambda url: calls.append(("arxiv", url)) or ExtractedPage(url=url, raw_content="arxiv", content_type="arxiv"),
    )
    monkeypatch.setattr(
        extraction_fetcher,
        "extract_pdf",
        lambda url, **kwargs: calls.append(("pdf", url, kwargs)) or ExtractedPage(url=url, raw_content="pdf", content_type="pdf"),
    )
    monkeypatch.setattr(
        extraction_fetcher,
        "extract_html",
        lambda url, **kwargs: calls.append(("html", url, kwargs)) or ExtractedPage(url=url, raw_content="html", content_type="html"),
    )

    pages = extraction_fetcher.fetch_many(
        [
            "https://arxiv.org/abs/2401.12345",
            "https://arxiv.org/pdf/2401.12345",
            "https://example.com/report.pdf",
            "https://example.com/page",
            "https://example.com/page#section",
        ],
        max_urls=10,
        max_chars_per_page=111,
        max_pdf_pages=2,
    )

    assert [page.content_type for page in pages] == ["arxiv", "arxiv", "pdf", "html"]
    assert calls[0] == ("arxiv", "https://arxiv.org/abs/2401.12345")
    assert calls[1] == ("arxiv", "https://arxiv.org/pdf/2401.12345")
    assert calls[2][0:2] == ("pdf", "https://example.com/report.pdf")
    assert calls[2][2]["max_chars_per_page"] == 111
    assert calls[2][2]["max_pages"] == 2
    assert calls[3][0:2] == ("html", "https://example.com/page")


def test_extraction_fetcher_enriches_search_items(monkeypatch):
    monkeypatch.setattr(
        extraction_fetcher,
        "fetch_many",
        lambda items, **kwargs: [
            ExtractedPage(
                url="https://example.com/a",
                title="Fetched title",
                raw_content="Full page text",
                content_type="html",
            )
        ],
    )

    enriched = extraction_fetcher.enrich_items_with_extracted_content(
        [
            {"url": "https://example.com/a#intro", "content": "Snippet"},
            {"url": "https://example.com/b", "content": "Untouched"},
        ]
    )

    assert enriched[0]["title"] == "Fetched title"
    assert enriched[0]["raw_content"] == "Full page text"
    assert enriched[0]["content"] == "Snippet"
    assert enriched[0]["extraction_status"] == "success"
    assert enriched[1] == {"url": "https://example.com/b", "content": "Untouched"}


@pytest.mark.skipif(
    os.getenv("ANNA_RESEARCHER_REAL_DDGS") != "1",
    reason="set ANNA_RESEARCHER_REAL_DDGS=1 to run the real DuckDuckGo/ddgs network test",
)
def test_duckduckgo_native_search_real_network():
    results = duckduckgo_native.search_duckduckgo("Anna app research", max_results=3)

    assert 1 <= len(results) <= 3
    for item in results:
        assert item["query"] == "Anna app research"
        assert item["url"].startswith(("http://", "https://"))
        assert item["title"] or item["content"]


def _native_definition():
    return {
        "id": "duckduckresearch",
        "name": "DuckDuckResearch",
        "native": {"adapter": "ddgs", "max_results": 5, "region": "wt-wt"},
    }


def _fixed_clock():
    values = iter([10.0, 10.125, 20.0, 20.125])
    return lambda: next(values)


# ---------------------------------------------------------------------------
# Envelope validation (ADR 0004)
# ---------------------------------------------------------------------------


def _user_envelope(**overrides):
    base = {
        "id": "custom",
        "name": "Custom",
        "request": {
            "method": "GET",
            "url": "https://api.example/search?token={token}&q={query}",
        },
        "pagination": {"mode": "none", "max_pages": 1},
        "result": {
            "items_path": "results[]",
            "url": {"mode": "path", "value": "url"},
            "title": {"mode": "path", "value": "title"},
            "content": {"mode": "paths", "value": ["snippet"]},
        },
        "response": {"content_type": "application/json"},
    }
    for key, value in overrides.items():
        if value is None:
            base.pop(key, None)
        else:
            base[key] = value
    return base


def test_envelope_accepts_minimal_user_definition():
    validate_envelope(_user_envelope(), kind="user")


@pytest.mark.parametrize(
    "mutator, expected_reason",
    [
        (lambda d: d["request"].__setitem__("method", "PUT"), "method_must_be_get_or_post"),
        (lambda d: d.__setitem__("auth", {"oauth_client_secret": "x"}), "oauth_not_supported"),
        (lambda d: d.__setitem__("auth", {"x-hmac-signature": "x"}), "hmac_not_supported"),
        (
            lambda d: d["request"].__setitem__("headers", {"Content-Type": "multipart/form-data; boundary=---"}),
            "content_type_not_supported",
        ),
        (lambda d: d.__setitem__("script", "() => {}"), "script_fields_not_supported"),
        (
            lambda d: d["request"].__setitem__("url", "https://api.example/search?token={token}&q={query}&fancy={magic}"),
            "unknown_placeholder",
        ),
        (
            lambda d: d["request"].__setitem__("url", "https://api.example/search?q={query}"),
            "token_placeholder_required",
        ),
        (lambda d: d.__setitem__("pagination", {"mode": "rolling", "max_pages": 1}), "pagination_mode_invalid"),
        (lambda d: d.__setitem__("pagination", {"mode": "page", "max_pages": 99}), "max_pages_exceeds_cap"),
        (lambda d: d["result"].pop("url"), "result_url_required"),
        (lambda d: d["result"].__setitem__("content", {"mode": "paths", "value": []}), "result_content_paths_must_be_nonempty_array"),
        (
            lambda d: d["result"].__setitem__("title", {"mode": "none"}),
            "result_title_mode_invalid",
        ),
        (
            lambda d: d["result"].__setitem__("url", {"mode": "template", "value": "https://example.test/{{context.token}}"}),
            "result_template_token_not_allowed",
        ),
        (
            lambda d: d["result"].__setitem__("content", {"mode": "template", "value": "bad {{query}}"}),
            "result_template_placeholder_invalid",
        ),
        (
            lambda d: d.__setitem__("response", {"content_type": "text/html"}),
            "response_must_be_json",
        ),
        (lambda d: d.__setitem__("max_parallel", 99), "max_parallel_out_of_range"),
    ],
)
def test_envelope_rejects_invalid_shapes(mutator, expected_reason):
    definition = _user_envelope()
    mutator(definition)
    with pytest.raises(EnvelopeError) as exc:
        validate_envelope(definition, kind="user")
    assert exc.value.reason == expected_reason


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------


def test_registry_lists_builtin_tavily_with_credential(tmp_path):
    root = tmp_path / ".research"
    creds = CredentialStore(root)
    registry = ResearchSourceRegistry(root, credentials=creds)
    views = registry.list_views()
    by_id = {view["id"]: view for view in views}
    assert list(by_id) == ["tavily", "duckduckgo"]
    assert by_id["tavily"]["kind"] == "builtin"
    assert by_id["tavily"]["credential_status"] == "missing"
    assert by_id["tavily"]["definition"]["id"] == "tavily"
    assert by_id["tavily"]["definition"]["request"]["body"]["api_key"] == "{token}"
    assert "credential" not in by_id["tavily"]["definition"]
    assert "token" not in by_id["tavily"]["definition"]
    assert by_id["duckduckgo"]["kind"] == "builtin"
    assert by_id["duckduckgo"]["credential_status"] == "configured"
    assert by_id["duckduckgo"]["credential"] == ""
    assert by_id["duckduckgo"]["definition"]["native"]["adapter"] == "ddgs"

    creds.set_token("tavily", "tvly-secret-1234")
    refreshed = registry.get_view("tavily")
    assert refreshed["credential_status"] == "configured"
    assert refreshed["credential"] == "tvly-secret-1234"


def test_registry_rejects_user_attempt_to_override_builtin(tmp_path):
    root = tmp_path / ".research"
    creds = CredentialStore(root)
    registry = ResearchSourceRegistry(root, credentials=creds)
    assert "tavily" in BUILTIN_SOURCE_IDS
    assert "duckduckgo" in BUILTIN_SOURCE_IDS
    with pytest.raises(ValidationError) as exc:
        registry.upsert_user_source(_user_envelope(id="tavily", name="hijacked"))
    assert exc.value.data.get("reason") == "builtin_protected"


def test_registry_upsert_and_delete_user_source_clears_credential(tmp_path):
    root = tmp_path / ".research"
    creds = CredentialStore(root)
    registry = ResearchSourceRegistry(root, credentials=creds)
    view = registry.upsert_user_source(_user_envelope(id="acme", name="ACME"))
    assert view["id"] == "acme"
    assert view["kind"] == "user"
    creds.set_token("acme", "acme-token-abcd")
    assert creds.get_token("acme") == "acme-token-abcd"
    registry.delete_user_source("acme")
    assert creds.get_token("acme") == ""
    with pytest.raises(NotFoundError):
        registry.get_definition("acme")


# ---------------------------------------------------------------------------
# Executor
# ---------------------------------------------------------------------------


class FakeResponse:
    def __init__(self, body: bytes, status: int = 200):
        self._body = body
        self.status = status

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False

    def read(self) -> bytes:
        return self._body


def fake_http(payload: dict, *, status: int = 200):
    body = json.dumps(payload).encode("utf-8")

    def _open(request, timeout=None):
        return FakeResponse(body, status=status)

    return _open


def test_executor_calls_builtin_tavily_with_token_substitution():
    captured: list[dict] = []

    def http(request, timeout=None):
        captured.append(
            {
                "method": request.get_method(),
                "url": request.full_url,
                "body": json.loads(request.data.decode("utf-8")) if request.data else None,
            }
        )
        return FakeResponse(
            json.dumps(
                {
                    "results": [
                        {"url": "https://ex.com/a", "title": "A", "content": "alpha"},
                        {"url": "https://ex.com/b", "title": "B", "content": "beta"},
                    ]
                }
            ).encode("utf-8")
        )

    executor = ResearchSourceExecutor(token_provider=lambda sid: "tvly-secret-abcd", http_open=http, sleep=lambda _: None)
    result = executor.call(builtin_tavily_definition(), "anna")
    assert result.error is None
    assert len(result.items) == 2
    assert result.items[0]["source_id"] == "tavily"
    assert captured and captured[0]["method"] == "POST"
    assert captured[0]["body"]["api_key"] == "tvly-secret-abcd"
    assert captured[0]["body"]["query"] == "anna"


def test_executor_returns_empty_result_when_no_items_returned():
    executor = ResearchSourceExecutor(
        token_provider=lambda sid: "tok",
        http_open=fake_http({"results": []}),
        sleep=lambda _: None,
    )
    result = executor.call(builtin_tavily_definition(), "anna")
    assert result.error == "empty_result"
    assert result.items == []


def test_executor_classifies_http_status_codes():
    cases = {
        401: "auth_failed",
        403: "auth_failed",
        429: "rate_limited",
        500: "upstream_5xx",
        502: "upstream_5xx",
        418: "bad_definition",
    }
    for status, expected in cases.items():
        def opener(status_code=status):
            def _open(request, timeout=None):
                raise urllib.error.HTTPError(request.full_url, status_code, "boom", hdrs=None, fp=io.BytesIO(b""))

            return _open

        executor = ResearchSourceExecutor(token_provider=lambda sid: "tok", http_open=opener(), sleep=lambda _: None)
        result = executor.call(builtin_tavily_definition(), "anna")
        assert result.error == expected, f"status {status} -> expected {expected}, got {result.error}"


def test_executor_retries_get_once_on_rate_limit():
    calls = {"n": 0}

    def opener(request, timeout=None):
        calls["n"] += 1
        if calls["n"] == 1:
            raise urllib.error.HTTPError(request.full_url, 429, "slow down", hdrs=None, fp=io.BytesIO(b""))
        return FakeResponse(json.dumps({"results": [{"url": "https://x", "title": "t", "snippet": "s"}]}).encode("utf-8"))

    sleeps: list[float] = []
    definition = _user_envelope()  # GET-method envelope
    executor = ResearchSourceExecutor(
        token_provider=lambda sid: "tok",
        http_open=opener,
        sleep=lambda d: sleeps.append(d),
    )
    result = executor.call(definition, "anna")
    assert calls["n"] == 2
    assert sleeps == [1.0]
    assert result.error is None
    assert len(result.items) == 1


def test_executor_supports_result_templates_and_single_object_items():
    definition = _user_envelope(
        result={
            "items_path": "company",
            "url": {"mode": "template", "value": "https://example.test/search?q={{context.query}}"},
            "title": {"mode": "template", "value": "{{item.name}} company profile"},
            "content": {
                "mode": "template",
                "value": "Company: {{item.name}}\nLegal person: {{item.people[0].name}}\nScope: {{item.scope}}",
            },
        }
    )
    executor = ResearchSourceExecutor(
        token_provider=lambda sid: "tok",
        http_open=fake_http({"company": {"name": "ACME", "people": [{"name": "Ada"}], "scope": "Research apps"}}),
        sleep=lambda _: None,
    )
    result = executor.call(definition, "anna app")
    assert result.error is None
    assert len(result.items) == 1
    assert result.items[0]["url"] == "https://example.test/search?q=anna app"
    assert result.items[0]["title"] == "ACME company profile"
    assert "Legal person: Ada" in result.items[0]["content"]


def test_executor_supports_url_none_and_rejects_scalar_items_path():
    none_url = _user_envelope(
        result={
            "items_path": "result",
            "url": {"mode": "none"},
            "title": {"mode": "path", "value": "name"},
            "content": {"mode": "paths", "value": ["scope"]},
        }
    )
    executor = ResearchSourceExecutor(
        token_provider=lambda sid: "tok",
        http_open=fake_http({"result": {"name": "ACME", "scope": "Research apps"}}),
        sleep=lambda _: None,
    )
    result = executor.call(none_url, "anna")
    assert result.error is None
    assert result.items[0]["url"] == ""

    scalar = dict(none_url)
    executor = ResearchSourceExecutor(
        token_provider=lambda sid: "tok",
        http_open=fake_http({"result": "not an object"}),
        sleep=lambda _: None,
    )
    result = executor.call(scalar, "anna")
    assert result.error == "bad_definition"


def test_executor_does_not_retry_post():
    calls = {"n": 0}

    def opener(request, timeout=None):
        calls["n"] += 1
        raise urllib.error.HTTPError(request.full_url, 429, "slow down", hdrs=None, fp=io.BytesIO(b""))

    executor = ResearchSourceExecutor(
        token_provider=lambda sid: "tok",
        http_open=opener,
        sleep=lambda _: None,
    )
    result = executor.call(builtin_tavily_definition(), "anna")
    assert calls["n"] == 1
    assert result.error == "rate_limited"


def test_executor_paginates_in_page_mode_until_empty():
    pages = [
        {"results": [{"url": "https://x.com/1", "title": "a", "content": "1"}]},
        {"results": [{"url": "https://x.com/2", "title": "b", "content": "2"}]},
        {"results": []},
    ]
    seen_pages: list[str] = []

    def opener(request, timeout=None):
        url = request.full_url
        seen_pages.append(url)
        payload = pages[len(seen_pages) - 1]
        return FakeResponse(json.dumps(payload).encode("utf-8"))

    definition = {
        "id": "pager",
        "name": "Pager",
        "request": {"method": "GET", "url": "https://x.com/?token={token}&q={query}&page={page}"},
        "pagination": {"mode": "page", "max_pages": 5, "page_size": 1, "start_page": 1},
        "result": {
            "items_path": "results[]",
            "url": {"mode": "path", "value": "url"},
            "title": {"mode": "path", "value": "title"},
            "content": {"mode": "paths", "value": ["content"]},
        },
        "response": {"content_type": "application/json"},
    }
    executor = ResearchSourceExecutor(token_provider=lambda sid: "tok", http_open=opener, sleep=lambda _: None)
    result = executor.call(definition, "anna")
    assert result.error is None
    assert len(result.items) == 2
    assert "page=1" in seen_pages[0]
    assert "page=2" in seen_pages[1]


def test_executor_test_returns_request_response_and_extracted_items():
    definition = builtin_tavily_definition()
    executor = ResearchSourceExecutor(
        token_provider=lambda sid: "tvly-secret-abcd",
        http_open=fake_http({"results": [{"url": "https://x.com/1", "title": "Title", "content": "Evidence"}]}),
        sleep=lambda _: None,
    )
    result = executor.test(definition, "anna")
    assert result.error is None
    assert result.pages[0]["request"]["method"] == "POST"
    assert result.pages[0]["request"]["body"]["api_key"] == "tvly-secret-abcd"
    assert result.pages[0]["response"]["json"]["results"][0]["title"] == "Title"
    assert result.extracted[0]["url"] == "https://x.com/1"
    assert result.extracted[0]["content"] == "Evidence"


def test_resolve_path_handles_dot_and_index_segments():
    payload = {"data": {"results": [{"name": "X"}, {"name": "Y"}]}}
    assert resolve_path(payload, "data.results[]") == payload["data"]["results"]
    assert resolve_path(payload, "data.results[0].name") == "X"
    assert resolve_path(payload, "data.missing") is None


def test_source_call_error_falls_back_to_bad_definition_for_unknown_codes():
    error = SourceCallError("nonsense", "boom")
    assert error.code == "bad_definition"


def test_call_section_research_source_uses_native_duckduckgo(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    dispatcher.native_executor = NativeResearchSourceExecutor(
        adapters={
            "ddgs": lambda query: [
                {"query": query, "url": "https://duck.example/section", "title": "Section result", "content": "section snippet"},
            ]
        },
        extractor=lambda items, **kwargs: [
            {**item, "icon": "https://duck.example/icon.png", "raw_content": "Section full extracted document"}
            for item in items
        ],
    )
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    dispatcher.dispatch(
        "app_save_confirmed_research_outline",
        {
            "research_id": job["research_id"],
            "sections": [
                {
                    "id": "section-1",
                    "title": "Background",
                    "outline": "Find background information.",
                    "allowed_source_ids": ["duckduckgo"],
                    "max_iterations": 1,
                }
            ],
        },
    )

    response = dispatcher.dispatch(
        "app_call_section_research_source",
        {
            "research_id": job["research_id"],
            "section_id": "section-1",
            "iteration": 1,
            "source_id": "duckduckgo",
            "queries": ["anna background"],
            "research_decision": {
                "type": "call_source",
                "knowledge_gap": "Missing background evidence",
                "rationale": "Needed for the section",
                "target_facet_ids": ["f1"],
            },
        },
    )

    call = response["source_call"]
    assert call["section_id"] == "section-1"
    assert call["source_id"] == "duckduckgo"
    assert call["results_count"] == 1
    loaded = dispatcher.jobs.load(job["research_id"])
    stored_call = loaded["section_iterations"]["section-1"][0]["source_calls"][0]
    assert "items" not in stored_call
    assert stored_call["results_count"] == 1
    assert stored_call["top_titles"] == ["Section result"]
    assert loaded["section_iterations"]["section-1"][0]["research_decision"] == {
        "type": "call_source",
        "knowledge_gap": "Missing background evidence",
        "rationale": "Needed for the section",
        "target_facet_ids": ["f1"],
    }
    item = loaded["section_iterations"]["section-1"][0]["raw_results"][0]
    assert item["source_name"] == "DuckDuckGo"
    assert item["url"] == "https://duck.example/section"
    assert item["content"] == "section snippet"
    assert item["icon"] == "https://duck.example/icon.png"
    assert item["document_id"]
    assert "url_body" not in item
    assert "raw_content" not in item
    document = dispatcher.web_documents.get(job["research_id"], item["document_id"])
    assert document and document["content"] == "Section full extracted document"
    assert document["content_type"] == "unknown"
    assert loaded["source_urls"] == []
    assert loaded["source_count"] == 0

    dispatcher.jobs.save_section_result(
        job["research_id"],
        "section-1",
        {
            "section_markdown": "## Background\n\nSelected evidence [1].",
            "section_summary": "Selected evidence.",
            "source_urls": ["https://duck.example/selected"],
            "status": "completed",
        },
    )
    loaded = dispatcher.jobs.load(job["research_id"])
    assert loaded["source_urls"] == ["https://duck.example/selected"]
    assert loaded["source_count"] == 1


def test_call_section_research_source_skips_duplicate_query(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    dispatcher.native_executor = NativeResearchSourceExecutor(
        adapters={
            "ddgs": lambda query: [
                {"query": query, "url": "https://duck.example/section", "title": "Section result", "content": "section snippet"},
            ]
        },
        extractor=lambda items, **kwargs: items,
    )
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    dispatcher.dispatch(
        "app_save_confirmed_research_outline",
        {
            "research_id": job["research_id"],
            "sections": [
                {
                    "id": "section-1",
                    "title": "Background",
                    "outline": "Find background information.",
                    "allowed_source_ids": ["duckduckgo"],
                    "max_iterations": 1,
                }
            ],
        },
    )
    dispatcher.dispatch(
        "app_call_section_research_source",
        {
            "research_id": job["research_id"],
            "section_id": "section-1",
            "iteration": 1,
            "source_id": "duckduckgo",
            "queries": ["anna background"],
        },
    )

    response = dispatcher.dispatch(
        "app_call_section_research_source",
        {
            "research_id": job["research_id"],
            "section_id": "section-1",
            "iteration": 2,
            "source_id": "duckduckgo",
            "queries": ["Anna  Background"],
        },
    )

    assert response["source_call"]["queries"] == []
    assert response["source_call"]["skipped_queries"] == ["Anna  Background"]
    assert response["source_call"]["results_count"] == 0
    assert response["source_call"]["error"] is None


def test_app_search_web_is_removed(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    job = dispatcher.dispatch("app_create_research_job", {"query": "anna"})["job"]
    with pytest.raises(ValidationError):
        dispatcher.dispatch("app_search_web", {"research_id": job["research_id"], "search_queries": ["anna"]})


def test_app_test_research_source_uses_draft_definition_and_saved_credential(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    dispatcher.dispatch("app_update_research_source_credential", {"id": "tavily", "credential": "tvly-secret-abcd"})

    def fake_http(request, timeout=None):
        return FakeResponse(
            json.dumps({"items": [{"href": "https://draft.example/a", "name": "Draft title", "body": "Draft body"}]}).encode("utf-8")
        )

    dispatcher.executor = ResearchSourceExecutor(
        token_provider=dispatcher._token_for, http_open=fake_http, sleep=lambda _: None
    )
    draft = dict(builtin_tavily_definition())
    draft["result"] = {
        "items_path": "items[]",
        "url": {"mode": "path", "value": "href"},
        "title": {"mode": "path", "value": "name"},
        "content": {"mode": "paths", "value": ["body"]},
    }
    immediate = dispatcher.dispatch(
        "app_test_research_source",
        {"id": "tavily", "definition": draft, "query": "anna"},
    )
    assert "test" not in immediate
    assert "Draft body" not in json.dumps(immediate)
    transfer = immediate["test_transfer"]
    assert transfer["method"] == "GET"
    result = get_json(transfer["url"])["test"]
    assert result["pages"][0]["request"]["body"]["api_key"] == "tvly-secret-abcd"
    assert result["pages"][0]["response"]["json"]["items"][0]["name"] == "Draft title"
    assert result["extracted"][0]["url"] == "https://draft.example/a"
    assert result["extracted"][0]["title"] == "Draft title"


def test_app_test_research_source_uses_native_definition_without_credential(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    dispatcher.native_executor = NativeResearchSourceExecutor(
        adapters={
            "ddgs": lambda query: [
                {"query": query, "url": "https://duck.example/test", "title": "Duck test", "content": "duck test body"},
            ]
        },
        extractor=lambda items, **kwargs: items,
    )

    immediate = dispatcher.dispatch(
        "app_test_research_source",
        {"id": "duckduckgo", "definition": builtin_duckduckgo_definition(), "query": "anna"},
    )

    result = get_json(immediate["test_transfer"]["url"])["test"]
    assert result["source_id"] == "duckduckgo"
    assert result["error"] is None
    assert result["pages"] == []
    assert result["extracted"][0]["url"] == "https://duck.example/test"


# ---------------------------------------------------------------------------
# Source list, credential updates, enabled flag
# ---------------------------------------------------------------------------


def test_app_list_research_sources_returns_builtin(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    sources = dispatcher.dispatch("app_list_research_sources", {})["sources"]
    by_id = {source["id"]: source for source in sources}
    assert list(by_id) == ["tavily", "duckduckgo"]
    assert by_id["tavily"]["kind"] == "builtin"
    assert by_id["tavily"]["credential_status"] == "missing"
    assert by_id["duckduckgo"]["kind"] == "builtin"
    assert by_id["duckduckgo"]["credential_status"] == "configured"


def test_update_research_source_credential_returns_plain_credential_and_clears(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    saved = dispatcher.dispatch(
        "app_update_research_source_credential", {"id": "tavily", "credential": "tvly-secret-abcd"}
    )["source"]
    assert saved["credential_status"] == "configured"
    assert saved["credential"] == "tvly-secret-abcd"
    cleared = dispatcher.dispatch("app_update_research_source_credential", {"id": "tavily", "clear": True})["source"]
    assert cleared["credential_status"] == "missing"
    assert cleared["credential"] == ""


def test_set_research_source_enabled_round_trip(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    disabled = dispatcher.dispatch("app_set_research_source_enabled", {"id": "tavily", "enabled": False})["source"]
    assert disabled["enabled"] is False
    enabled = dispatcher.dispatch("app_set_research_source_enabled", {"id": "tavily", "enabled": True})["source"]
    assert enabled["enabled"] is True


def test_upsert_user_source_persists_and_credential_can_be_added(tmp_path):
    dispatcher = make_dispatcher(tmp_path)
    definition = _user_envelope(id="acme", name="ACME")
    upserted = dispatcher.dispatch(
        "app_upsert_research_source",
        {"definition": definition, "credential": "acme-token-xyz1"},
    )["source"]
    assert upserted["kind"] == "user"
    assert upserted["credential_status"] == "configured"
    listed = dispatcher.dispatch("app_list_research_sources", {})["sources"]
    assert any(s["id"] == "acme" for s in listed)
    dispatcher.dispatch("app_delete_research_source", {"id": "acme"})
    listed_after = dispatcher.dispatch("app_list_research_sources", {})["sources"]
    assert all(s["id"] != "acme" for s in listed_after)


# ---------------------------------------------------------------------------
# Context selector: source-prefixed emission and URL-empty fallback
# ---------------------------------------------------------------------------


def test_selector_emits_source_prefix_and_dedupes_by_url():
    selector = LexicalContextSelector(max_sources=3, max_per_domain=2, context_budget=2000)
    long_context = "Anna app research context " * 16
    long_details = "anna app research details " * 16
    selected = selector.select(
        query="anna app research",
        search_queries=["anna app research"],
        search_results=[
            {"query": "anna", "source_id": "tavily", "source_name": "Tavily", "url": "https://example.com/a", "title": "Anna research", "content": long_context},
            {"query": "anna", "source_id": "tavily", "source_name": "Tavily", "url": "https://example.com/a", "title": "Duplicate", "content": "duplicate"},
            {"query": "anna", "source_id": "acme", "source_name": "ACME", "url": "", "title": "Same title", "content": long_details},
            {"query": "anna", "source_id": "acme", "source_name": "ACME", "url": "", "title": "Same title", "content": "different body"},
        ],
    )
    text = selected["selected_context"]
    assert "[来源: Tavily]" in text
    assert "[来源: ACME]" in text
    # url-empty entries dedupe on (source_id, title)
    assert selected["source_urls"] == ["https://example.com/a"]
    assert text.count("Same title") == 1


def test_selector_skips_failed_and_short_extractions():
    selector = LexicalContextSelector(max_sources=3, max_per_domain=3, context_budget=2000, min_content_length=120)
    selected = selector.select(
        query="starbucks china market",
        search_queries=["starbucks china market"],
        search_results=[
            {
                "query": "starbucks",
                "source_id": "duckduckgo",
                "source_name": "DuckDuckGo",
                "url": "https://example.com/failed",
                "title": "Starbucks China",
                "content": "Starbucks China market " * 20,
                "extraction_status": "failed",
            },
            {
                "query": "starbucks",
                "source_id": "duckduckgo",
                "source_name": "DuckDuckGo",
                "url": "https://example.com/short",
                "title": "Starbucks China",
                "content": "Starbucks China market",
            },
            {
                "query": "starbucks",
                "source_id": "duckduckgo",
                "source_name": "DuckDuckGo",
                "url": "https://example.com/good",
                "title": "Starbucks China market report",
                "content": "Starbucks China market expansion and competition analysis " * 12,
            },
        ],
    )

    assert selected["source_urls"] == ["https://example.com/good"]
    assert "failed" not in selected["selected_context"]
    assert "short" not in selected["selected_context"]


class KeywordEmbeddings:
    def __init__(self, *, fail=False):
        self.fail = fail
        self.batch_counts = []

    def create_batches(self, *, batches, model="anna-managed-v1", timeout=30.0):
        if self.fail:
            raise EmbeddingsError(-32505, "forced embedding failure")
        self.batch_counts.append(len(batches))
        return [
            {
                "data": [
                    {
                        "embedding": [
                            float(str(text).casefold().count("alpha")),
                            float(str(text).casefold().count("beta")),
                            1.0,
                        ]
                    }
                    for text in batch
                ],
                "_meta": {"dimensions": 3},
            }
            for batch in batches
        ]


def test_hybrid_selector_chunks_rrf_groups_and_does_not_persist_vectors(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="alpha beta")
    documents = WebDocumentStore(jobs)
    alpha_text = ("alpha evidence and supporting detail. " * 300).strip()
    beta_text = ("beta evidence and supporting detail. " * 300).strip()
    alpha_id = documents.put_page(job["research_id"], ExtractedPage(url="https://example.com/alpha", title="Alpha", raw_content=alpha_text, content_type="html"))
    beta_id = documents.put_page(job["research_id"], ExtractedPage(url="https://example.org/beta", title="Beta", raw_content=beta_text, content_type="html"))
    embeddings = KeywordEmbeddings()
    selector = HybridContextSelector(embeddings=embeddings, documents=documents)

    selected = selector.select(
        research_id=job["research_id"],
        query="alpha beta",
        search_queries=["alpha evidence", "beta evidence"],
        search_results=[
            {"document_id": alpha_id, "url": "https://example.com/alpha", "title": "Alpha", "source_id": "duckduckgo", "source_name": "DuckDuckGo", "query": "alpha evidence"},
            {"document_id": beta_id, "url": "https://example.org/beta", "title": "Beta", "source_id": "duckduckgo", "source_name": "DuckDuckGo", "query": "beta evidence"},
            {"url": "https://tavily.example/alpha", "title": "Tavily alpha", "source_id": "tavily", "source_name": "Tavily", "query": "alpha evidence", "content": "alpha summary evidence"},
        ],
    )

    assert sum(len(source["selected_chunks"]) for source in selected["selected_sources"]) <= 16
    assert "[Chunk" in selected["selected_context"]
    assert any(source["source_id"] == "tavily" for source in selected["selected_sources"])
    assert all("rrf_score" in chunk for source in selected["selected_sources"] for chunk in source["selected_chunks"])
    assert all("matched_queries" in chunk for source in selected["selected_sources"] for chunk in source["selected_chunks"])
    assert max(embeddings.batch_counts) <= 8
    assert 8 in embeddings.batch_counts
    assert len(embeddings.batch_counts) >= 3
    persisted = "".join(path.read_text(encoding="utf-8") for path in (jobs.job_dir_for(job["research_id"]) / "web_documents").glob("*.json"))
    assert "embedding" not in persisted


class OrthogonalKeywordEmbeddings:
    def create_batches(self, *, batches, model="anna-managed-v1", timeout=30.0):
        return [
            {
                "data": [
                    {
                        "embedding": [
                            1.0 if "alpha" in str(text).casefold() else 0.0,
                            1.0 if "beta" in str(text).casefold() else 0.0,
                        ]
                    }
                    for text in batch
                ],
                "_meta": {"dimensions": 2},
            }
            for batch in batches
        ]


def test_hybrid_selector_keeps_top_eight_independently_for_each_query(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="alpha beta")
    selector = HybridContextSelector(embeddings=OrthogonalKeywordEmbeddings(), documents=WebDocumentStore(jobs))
    results = [
        {
            "url": f"https://alpha.example/{index}",
            "title": f"Alpha {index}",
            "source_id": "tavily",
            "source_name": "Tavily",
            "content": f"alpha evidence item {index}",
        }
        for index in range(10)
    ] + [
        {
            "url": f"https://beta.example/{index}",
            "title": f"Beta {index}",
            "source_id": "tavily",
            "source_name": "Tavily",
            "content": f"beta evidence item {index}",
        }
        for index in range(10)
    ] + [{
        "url": "https://neutral.example/item",
        "title": "Neutral",
        "source_id": "tavily",
        "source_name": "Tavily",
        "content": "unrelated neutral material",
    }]

    selected = selector.select(
        research_id=job["research_id"],
        query="alpha beta",
        search_queries=["alpha", "beta"],
        search_results=results,
    )

    chunks = [chunk for source in selected["selected_sources"] for chunk in source["selected_chunks"]]
    assert len(chunks) == 16
    assert sum(any(match["query"] == "alpha" for match in chunk["matched_queries"]) for chunk in chunks) == 8
    assert sum(any(match["query"] == "beta" for match in chunk["matched_queries"]) for chunk in chunks) == 8
    assert max(chunk["rrf_score"] for chunk in chunks) == pytest.approx(2 / 61, abs=1e-8)


def test_hybrid_selector_applies_embedding_similarity_threshold(tmp_path):
    class ThresholdEmbeddings:
        def create_batches(self, *, batches, model="anna-managed-v1", timeout=30.0):
            vectors = []
            for batch in batches:
                data = []
                for text in batch:
                    value = str(text).casefold()
                    vector = [1.0, 0.0] if value.strip() == "target" else [0.3, math.sqrt(0.91)]
                    data.append({"embedding": vector})
                vectors.append({"data": data, "_meta": {"dimensions": 2}})
            return vectors

    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="target")
    selector = HybridContextSelector(embeddings=ThresholdEmbeddings(), documents=WebDocumentStore(jobs))
    selected = selector.select(
        research_id=job["research_id"],
        query="target",
        search_queries=["target"],
        search_results=[{
            "url": "https://example.com/unrelated",
            "title": "Unrelated",
            "source_id": "tavily",
            "source_name": "Tavily",
            "content": "orthogonal material without the search term",
        }],
    )

    assert selected["selected_sources"] == []


def test_hybrid_selector_seed_diversity_prefers_new_source_per_facet(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="alpha beta gamma")
    selector = HybridContextSelector(embeddings=KeywordEmbeddings(), documents=WebDocumentStore(jobs))

    selected = selector.select(
        research_id=job["research_id"],
        query="alpha beta gamma",
        search_queries=["alpha", "beta", "gamma"],
        search_results=[
            {
                "url": "https://example.com/general",
                "title": "General",
                "source_id": "tavily",
                "source_name": "Tavily",
                "content": "alpha beta gamma " * 20,
            },
            {
                "url": "https://example.com/beta",
                "title": "Beta",
                "source_id": "tavily",
                "source_name": "Tavily",
                "content": "beta evidence " * 20,
            },
            {
                "url": "https://example.com/gamma",
                "title": "Gamma",
                "source_id": "tavily",
                "source_name": "Tavily",
                "content": "gamma evidence " * 20,
            },
        ],
        diversify_queries=True,
    )

    assert selected["source_urls"][:3] == [
        "https://example.com/general",
        "https://example.com/beta",
        "https://example.com/gamma",
    ]


def test_web_document_store_deduplicates_fragment_urls(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="dedupe")
    documents = WebDocumentStore(jobs)

    first = documents.put_page(job["research_id"], ExtractedPage(url="https://example.com/a#one", raw_content="first full body"))
    second = documents.put_page(job["research_id"], ExtractedPage(url="https://example.com/a#two", raw_content="second full body"))

    assert first == second
    files = list((jobs.job_dir_for(job["research_id"]) / "web_documents").glob("*.json"))
    assert len(files) == 2  # one document plus index.json
    assert documents.get(job["research_id"], first)["content"] == "second full body"


def test_recursive_chunk_ranges_use_1000_size_and_100_overlap():
    ranges = split_text_ranges("alpha sentence. " * 300)

    assert len(ranges) > 2
    assert all(end - start <= 1000 for start, end in ranges)
    assert all(current[0] == previous[1] - 100 for previous, current in zip(ranges, ranges[1:]))


def test_hybrid_selector_fails_without_bm25_fallback_on_embedding_error(tmp_path):
    jobs = JobStore(root=tmp_path / ".research")
    job = jobs.create(query="alpha")
    documents = WebDocumentStore(jobs)
    document_id = documents.put_page(job["research_id"], ExtractedPage(url="https://example.com/a", raw_content="alpha evidence " * 100))
    selector = HybridContextSelector(embeddings=KeywordEmbeddings(fail=True), documents=documents)

    with pytest.raises(EmbeddingsError):
        selector.select(
            research_id=job["research_id"],
            query="alpha",
            search_queries=["alpha"],
            search_results=[{"document_id": document_id, "url": "https://example.com/a", "content": "alpha", "source_id": "duckduckgo"}],
        )
